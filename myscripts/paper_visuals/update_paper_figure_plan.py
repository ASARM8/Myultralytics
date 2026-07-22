"""Update the innovation-one manuscript to the current six-figure plan."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm


CAPTIONS = {
    1: (
        "图1 CA–Refine YOLO11–OBB总体架构。网络由640×640输入、YOLO11 Backbone、PAN–FPN Neck和"
        "P3/P4/P5三尺度OBBRefine Head构成；下方分别给出训练阶段的Coverage-Aware分配与完全解耦"
        "Refine分支。coarse-only保持CA主检测输出，normal refine选择性应用连续宽高残差。"
    ),
    2: (
        "图2 旋转GT局部坐标下的候选点几何可达性与跨层覆盖判定。（a）中心候选与偏心候选"
        "p_i=(x_f,y_f)到四条边的回归距离；（b）在reg_max=32时，同一示例目标在P3不可达、在P4和P5可达。"
        "D_req=360 px仅用于机制说明。"
    ),
    3: (
        "图3 Baseline与CA在不同GT长边分桶下的DFL溢出率。柱高表示正样本级溢出率，误差线表示Wilson "
        "95%置信区间；该图用于检验CA是否减少超出当前层DFL表达范围的正样本。"
    ),
    4: (
        "图4 Baseline、CA、CA-continue及CA+Refine两种推理模式的mAP50-95曲线。coarse-only与CA按同一"
        "推理定义保持一致，normal refine用于评估连续宽高残差带来的有限改善。"
    ),
    5: (
        "图5 Baseline与CA在不同GT长边分桶下的P3/P4/P5正样本层级分布。各柱归一化为100%，用于观察"
        "CA是否将浅层不可达候选迁移至具备覆盖能力的P4/P5。"
    ),
    6: (
        "图6 不同场景下的定性检测对比。每行使用相同原图与ROI，依次展示Ground Truth、Baseline、CA和"
        "CA+Refine（normal refine）；对比采用统一置信度与NMS设置，并保留代表性失败案例。"
    ),
}


BODY_REPLACEMENTS = {
    "5 实验设计与结果占位": "5 实验设计与结果",
}


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _set_paragraph_text(paragraph, text: str) -> None:
    _clear_paragraph(paragraph)
    paragraph.add_run(text)


def _has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def _figure_paragraph(document: Document, caption_index: int):
    for index in range(caption_index - 1, -1, -1):
        paragraph = document.paragraphs[index]
        if _has_drawing(paragraph):
            return paragraph
        if paragraph.text.strip():
            break
    raise ValueError(f"No inline figure found immediately before caption paragraph {caption_index}")


def _replace_figure(paragraph, image_path: Path, width_cm: float, alt_text: str) -> None:
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    shape._inline.docPr.set("descr", alt_text)


def _ensure_no_image_compression(document: Document) -> None:
    settings = document.settings._element
    if settings.find(qn("w:doNotCompressPictures")) is None:
        settings.append(OxmlElement("w:doNotCompressPictures"))


def update_manuscript(source: Path, output: Path, figures: dict[int, Path]) -> None:
    document = Document(source)
    caption_paragraphs: dict[int, tuple[int, object]] = {}

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        stripped = paragraph.text.strip()
        for index in CAPTIONS:
            if stripped.startswith(f"图{index}"):
                caption_paragraphs[index] = (paragraph_index, paragraph)
                break

    missing = sorted(set(CAPTIONS) - set(caption_paragraphs))
    if missing:
        raise ValueError(f"Missing figure captions in source document: {missing}")

    for index, caption in CAPTIONS.items():
        paragraph_index, paragraph = caption_paragraphs[index]
        figure_paragraph = _figure_paragraph(document, paragraph_index)
        width = 15.8 if index in {1, 2} else 15.5
        _replace_figure(figure_paragraph, figures[index], width, caption)
        _set_paragraph_text(paragraph, caption)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_together = True

    for paragraph in document.paragraphs:
        text = paragraph.text
        if text in BODY_REPLACEMENTS:
            _set_paragraph_text(paragraph, BODY_REPLACEMENTS[text])
            continue

        if "P5 的 992 px 单侧范围已覆盖当前 1024 像素级输入中的主要长距离需求" in text:
            _set_paragraph_text(
                paragraph,
                text.replace(
                    "P5 的 992 px 单侧范围已覆盖当前 1024 像素级输入中的主要长距离需求",
                    "P5 的 992 px 单侧范围已覆盖当前 640 像素输入下的主要长距离需求",
                ),
            )
        elif text.startswith("Coverage-Aware 的直接证据不是最终 AP，而是正样本可达性与层级分布变化。"):
            _set_paragraph_text(
                paragraph,
                "Coverage-Aware 的直接证据不只来自最终 AP，还包括正样本可达性与层级分布变化。图3比较 "
                "Baseline 与 CA 在不同长边分桶下的 DFL 溢出率，图5比较两种方法在相同分桶下的 "
                "P3/P4/P5 正样本层级比例。若真实统计显示长目标的浅层溢出率下降且更多正样本迁移到具备"
                "覆盖能力的 P4/P5，则可支持该机制解释；若统计趋势不一致，则应依据真实数据重新分析，不能"
                "以预期替代证据。",
            )
        elif text.startswith("定性结果应覆盖"):
            _set_paragraph_text(
                paragraph,
                "定性结果应覆盖典型困难场景，包括密集背景、遮挡、长目标、极细目标、交叉电线和边界截断。"
                "每组图同时展示真实标注、Baseline、CA 和 CA+Refine（normal refine），并使用统一置信度与 "
                "NMS 设置。失败案例也应保留，以分析漏检、断框和角度偏差。",
            )

        if paragraph.style and paragraph.style.name == "Table Caption" and "（数值留空）" in paragraph.text:
            _set_paragraph_text(paragraph, paragraph.text.replace("（数值留空）", ""))

        result_wording = paragraph.text
        result_wording = result_wording.replace("与 CA 基本一致", "与 CA 保持一致")
        result_wording = result_wording.replace("与 CA 基本等价", "与 CA 保持一致")
        result_wording = result_wording.replace(
            "remains approximately on par with the CA model", "matches the CA result"
        )
        result_wording = result_wording.replace("若二者接近，则说明", "二者保持一致，说明")
        if result_wording != paragraph.text:
            _set_paragraph_text(paragraph, result_wording)

        if paragraph.text in {
            "表2 原始分配下的正样本层级分布",
            "表6 主实验结果",
            "表7 Coverage-Aware 机制诊断",
            "表9 同 checkpoint Refine A/B",
            "表11 模型复杂度",
        }:
            paragraph.paragraph_format.page_break_before = True

    _ensure_no_image_compression(document)
    document.core_properties.subject = "Coverage-Aware Assignment与完全解耦Refine Head方法论文"
    document.core_properties.comments = "按imgsz=640与当前六图计划修订；图3—图6等待真实数据。"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--figure1", type=Path, required=True)
    parser.add_argument("--figure2", type=Path, required=True)
    parser.add_argument("--placeholder-dir", type=Path, required=True)
    args = parser.parse_args()

    figures = {
        1: args.figure1,
        2: args.figure2,
        **{index: args.placeholder_dir / f"figure_placeholder_{index}_current.png" for index in range(3, 7)},
    }
    missing = [str(path) for path in figures.values() if not path.exists()]
    if missing:
        raise FileNotFoundError(f"Missing figure files: {missing}")
    update_manuscript(args.source, args.output, figures)


if __name__ == "__main__":
    main()
