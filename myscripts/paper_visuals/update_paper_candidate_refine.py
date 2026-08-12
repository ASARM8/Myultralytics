"""Update the innovation-one paper to the frozen proposal-refine method and audited results.

This script preserves the existing thesis styles and equation/caption positions.  It
creates a new DOCX instead of overwriting the prior draft.  Native equations marked
here are converted by ``insert_word_equations_refine.ps1`` in a second, Word-COM step.
"""

from __future__ import annotations

import argparse
import csv
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--figure1", type=Path, required=True)
    parser.add_argument("--figure4", type=Path, required=True)
    parser.add_argument("--val-csv", type=Path, required=True)
    return parser.parse_args()


def read_variants(path: Path) -> dict[str, dict[str, float]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    result: dict[str, dict[str, float]] = {}
    for row in rows:
        result[row["variant"]] = {
            key: float(row[key])
            for key in ("precision", "recall", "map50", "map50_95", "ap75", "ap90", "ap95")
        }
    if not {"coarse", "identity", "refined"}.issubset(result):
        raise ValueError("validation CSV must contain coarse, identity and refined rows")
    return result


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def replace_text(paragraph, text: str, *, bold_prefix: str | None = None) -> None:
    old_rpr = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        old_rpr = deepcopy(paragraph.runs[0]._r.rPr)
    clear_paragraph(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        lead = paragraph.add_run(bold_prefix)
        lead.bold = True
        rest = paragraph.add_run(text[len(bold_prefix) :])
        if old_rpr is not None:
            rest._r.insert(0, deepcopy(old_rpr))
    else:
        run = paragraph.add_run(text)
        if old_rpr is not None:
            run._r.insert(0, deepcopy(old_rpr))


def set_title(paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def replace_picture(paragraph, image_path: Path, width_inches: float) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)


def set_cell(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text)


def remove_row(table, index: int) -> None:
    row = table.rows[index]
    table._tbl.remove(row._tr)


def remove_column(table, index: int) -> None:
    grid = table._tbl.tblGrid
    if grid is not None and len(grid.gridCol_lst) > index:
        grid.remove(grid.gridCol_lst[index])
    for row in table.rows:
        cells = row._tr.tc_lst
        if len(cells) > index:
            row._tr.remove(cells[index])


def fmt(value: float, sign: bool = False) -> str:
    return f"{value:+.4f}" if sign else f"{value:.4f}"


def main() -> None:
    args = parse_args()
    for path in (args.input, args.figure1, args.figure4, args.val_csv):
        if not path.is_file():
            raise FileNotFoundError(path)
    metrics = read_variants(args.val_csv)
    coarse, refined = metrics["coarse"], metrics["refined"]

    args.output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(args.input, args.output)
    doc = Document(args.output)
    p = doc.paragraphs

    set_title(p[0], "面向低空细长障碍物的覆盖能力感知与候选框级几何精修方法")
    replace_text(
        p[2],
        "摘要：低空巡检图像中的电线、拉线和细杆具有长边跨度大、短边像素少、方向任意及背景干扰强等特点。现有旋转目标检测器仍可能把长目标分配给回归范围不足的浅层候选点，并在高交并比条件下受到短边尺度误差的显著影响。针对上述问题，本文提出覆盖能力感知与候选框级几何精修方法。首先，在旋转框局部坐标系中计算候选点完整覆盖真实框所需的最大距离，并在任务对齐分配前过滤超出当前特征层分布式回归上限的候选点；其次，将 reg_max 扩展为 32，为重新分配后的中长目标提供足够的距离表示范围；最后，以 CA 检测结果为候选框，利用 P2/P3 局部特征进行方向对齐的旋转 ROI 采样，仅预测短边和长边的连续尺度残差，同时保持中心、角度、类别置信度及 NMS 结果不变。在固定的 FP32、batch=8、imgsz=640 验证口径下，候选框级精修将 CA 的 mAP50-95 从 0.4541 提升至 0.5625，AP75 从 0.4398 提升至 0.5588，AP90 从 0.2457 提升至 0.2614；AP95 由 0.0990 变为 0.0972。恒等路径与 CA 完全一致，匹配候选框的平均 IoU 增量为 0.0466，改善比例为 63.41%。结果表明，该精修路径在当前固定实验设置下具有明确的正向作用，但 AP95 的轻微下降和单随机种子设置仍要求对结论保持审慎。",
        bold_prefix="摘要：",
    )
    replace_text(p[3], "关键词：低空障碍物检测；电线检测；旋转目标检测；覆盖能力感知；候选框级精修；旋转 ROI")
    replace_text(
        p[5],
        "Power lines, guy wires, and thin poles in low-altitude inspection imagery have long spatial extents, extremely narrow short sides, arbitrary orientations, and weak contrast against cluttered backgrounds. Existing oriented detectors may assign elongated targets to shallow candidates whose distributional regression ranges are insufficient, while small short-side errors can strongly affect high-IoU localization. This paper presents a coverage-aware assignment and proposal-level geometric refinement method. Geometrically unreachable candidates are filtered before task-aligned assignment, and reg_max is enlarged to 32. The frozen CA detector then provides post-NMS coarse proposals. Direction-aligned local features are sampled from P2 and P3 by rotated ROI extraction, and an independent refiner predicts only continuous short- and long-side scale residuals while preserving center, angle, confidence, and NMS results. Under the fixed FP32, batch-8, 640-pixel validation protocol, refinement improves mAP50-95 from 0.4541 to 0.5625 and AP75 from 0.4398 to 0.5588. AP90 increases from 0.2457 to 0.2614, whereas AP95 slightly decreases from 0.0990 to 0.0972. The identity path exactly reproduces the CA output, and the mean IoU change over matched proposals is +0.0466. These results support a clear positive effect under the current controlled setting without implying statistical significance.",
    )
    replace_text(p[6], "Key words: low-altitude obstacle detection; power-line detection; oriented object detection; coverage-aware assignment; proposal refinement; rotated ROI")

    replace_text(
        p[16],
        "（4）针对 CA 输出中的剩余尺度误差，设计候选框级局部几何精修路径。该路径以 post-NMS coarse OBB 为候选框，从 P2/P3 提取方向对齐的旋转 ROI 特征，只学习短边和长边的连续尺度残差；训练时冻结 CA，推理时保持中心、角度、类别置信度和 NMS 结果不变，从而使精修作用可以通过同权重恒等对照直接归因。",
    )
    replace_picture(p[17], args.figure1, 6.25)
    replace_text(
        p[18],
        "图1 CA–Refine YOLO11–OBB 总体架构。上部为 640×640 输入、YOLO11 Backbone、PAN–FPN Neck 与三尺度 OBB Detect Head；下部给出训练阶段的 Coverage-Aware Assignment 以及以 post-NMS coarse OBB 和 P2/P3 特征为输入的候选框级局部几何精修。",
    )

    replace_text(
        p[58],
        "本文方法由覆盖能力感知分配、最小充分回归扩容和候选框级局部几何精修三部分组成。前两部分共同构成 CA 检测器：分配器约束候选点是否具备完整表达真实框的能力，reg_max=32 为可达层级提供足够的距离表示范围。第三部分不再附着于密集检测头，而是在 CA 完成解码与 NMS 后，对保留下来的 coarse OBB 逐个进行局部尺度校正。这样既保留总体“CA + reg_max + Refine”框架，又把精修对象限定为真实参与最终预测的候选框。",
    )
    replace_text(p[74], "4.4 候选框级局部几何精修")
    replace_text(
        p[75],
        "CA 检测器输出 post-NMS coarse 旋转框 [[M_BC]]。对于每个候选框，首先将宽、高转换为与 OBB 等价表示无关的短边 [[M_SC]] 和长边 [[M_LC]]：",
    )
    replace_text(p[76], "[[EQ7]]\t\t(7)")
    replace_text(
        p[77],
        "以 coarse OBB 的长轴方向建立局部坐标系，并从冻结的 P2、P3 特征中执行 5×24 旋转 ROI 采样。长轴上下文系数为 1.2，短轴上下文系数为 4.0，且短轴采样范围不小于 16 px。两个尺度的 ROI 特征经投影、融合卷积和 MLP 后，仅输出短边与长边的对数尺度残差。真实监督目标定义为：",
    )
    replace_text(p[78], "[[EQ8]]\t\t(8)")
    replace_text(
        p[79],
        "为避免极端样本经硬裁剪后集中在边界，目标使用分符号 tanh 平滑压缩，并只占用物理输出范围的 80%。短边物理残差范围为 [-0.50, 0.20]，长边范围为 [-0.08, 0.08]；监督范围相应保留 20% 输出余量。采用按 TAL 匹配权重归一化的 SmoothL1 损失（beta=0.05），并以短边像素尺度构造不低于 0.25 的小目标权重。其训练目标可概括为：",
    )
    replace_text(p[80], "[[EQ9]]\t\t(9)")
    replace_text(
        p[81],
        "精修训练采用 geometry-only 口径：CA 主干及检测头保持 eval 状态并冻结参数，coarse 框在进入精修路径前停止梯度，优化器只更新新增精修器。推理阶段对全部 post-NMS proposal 执行一次精修，不使用质量门控，也不进行第二次 NMS。中心、角度与类别置信度保持 coarse 输出，仅按短、长边残差更新尺度：",
    )
    replace_text(p[82], "[[EQ10]]\t\t(10)")
    replace_text(p[83], "4.5 训练、选择与验证口径")
    replace_text(
        p[84],
        "训练从固定 CA checkpoint 初始化，使用原训练集中的确定性 20% image-level holdout 进行检查点选择。精修器训练 15 个 epoch，每轮在 holdout 上评估，最终选取 epoch 4。正式结果不直接使用训练期 holdout 数值，而是在固定 CA 权重和选定精修权重后，以 FP32、batch=8、imgsz=640 在独立验证集重新运行 coarse、identity 和 refined 三种模式。",
    )
    replace_text(
        p[85],
        "coarse 模式完全绕过精修输出，identity 模式执行精修链路但强制零残差，refined 模式使用实际残差。coarse 与 identity 的七项指标完全一致，且 CA 权重哈希在精修训练前后保持不变，说明主检测链路未被修改。因而 refined 与 coarse 的差值可以归因于候选框级尺度校正。论文统一采用预先固定的 FP32、batch=8 验证口径，不把不同数值精度或 batch 设置混入主结果。",
    )

    replace_text(
        p[92],
        "基线采用 YOLO11l-OBB，CA 模型使用 Coverage-Aware Assignment 并设 reg_max=32，最终模型在固定 CA checkpoint 后连接候选框级几何精修器。精修器使用 P2/P3 特征、32 个 ROI 通道、5×24 旋转 ROI 和 128 维隐藏层，仅训练 15 个 epoch；优化器为 AdamW，初始学习率 3×10^-4，weight decay 为 1×10^-4，warmup 为 3 个 epoch，随机种子为 0。输入尺寸固定为 640，正式验证采用 FP32、batch=8。Baseline 与 CA 的其余训练设置由原始训练日志补入表5。",
    )
    replace_text(p[95], "主要指标采用 mAP50-95，辅助报告 Precision、Recall、mAP50、AP75、AP90 和 AP95。mAP50 衡量目标是否被大致检出，mAP50-95 与 AP75 更能反映边界贴合质量；AP90 和 AP95 用于观察极严格 IoU 阈值下的收益与代价。效率指标包括参数量、FLOPs、单图延迟和 FPS。")
    replace_text(
        p[97],
        "主结果以相同验证设置比较 Baseline、CA 与 CA+Refine。当前已完成并锁定的是 CA 与 CA+Refine 的同权重 FP32/batch=8 对照；Baseline 的统一独立验证结果仍按相同协议补采。coarse 和 identity 仅用于内部恒等性校验，不作为两个独立方法重复计入主表。",
    )
    replace_text(
        p[100],
        "在独立验证集上，CA 的 Precision、Recall、mAP50 和 mAP50-95 分别为 0.8044、0.7738、0.7395 和 0.4541；启用候选框级精修后分别达到 0.8939、0.8665、0.8936 和 0.5625，其中 mAP50-95 提高 0.1084。AP75 由 0.4398 提高至 0.5588，AP90 由 0.2457 提高至 0.2614；AP95 由 0.0990 轻微下降至 0.0972。该结果说明精修对中高 IoU 定位具有明确改善，但在极严格阈值下仍存在少量边界样本被负向调整的情况。由于当前正式结果来自单随机种子，本文表述为当前固定设置下的明确正向作用，不作统计显著性声明。",
    )
    replace_picture(p[101], args.figure4, 6.25)
    replace_text(
        p[102],
        "图4 候选框级几何精修的训练选择与独立验证结果。（a）15 个训练 epoch 在 holdout 上的 mAP50-95，epoch 4 被选为最终检查点；（b）固定 FP32、batch=8 验证口径下 CA 与 CA+Refine 的 mAP50-95、AP75、AP90 和 AP95 对比。",
    )

    replace_text(p[109], "5.5 高 IoU 几何误差与 Refine 恒等对照")
    replace_text(
        p[110],
        "除聚合 AP 外，本文从候选框匹配层面检查精修是否真正改善几何。固定验证中共有 5260 个有效 proposal，其中 3444 个与真实框匹配；匹配 proposal 的平均 IoU 增量为 0.0466，改善比例为 63.41%，高于 36.59% 的恶化比例。短边与长边残差的边界命中率均为 0，排除了输出被统一推向限制边界的退化模式。表8的 Oracle 几何替换仍作为后续细化不同自由度上限的补充实验。",
    )
    replace_text(p[113], "同权重恒等对照使用同一 CA checkpoint 和同一精修 checkpoint。coarse 模式绕过精修，identity 模式强制零残差，refined 模式启用实际短、长边残差；数据、阈值、NMS、数值精度和 batch 设置保持一致。")
    replace_text(
        p[116],
        "表9表明，coarse 与 identity 的各项指标完全一致，说明旋转 ROI、特征融合和结果回写本身不会改变 CA 输出。启用实际残差后，mAP50-95、AP75 和 AP90 分别提高 0.1084、0.1190 和 0.0157；AP95 下降 0.0018。结合平均 IoU 与改善比例，可判断主要收益来自实例相关的尺度校正，而不是固定缩放、分数重标定或二次 NMS。",
    )
    replace_text(
        p[122],
        "reg_max 扩容会增加边界分布输出通道，候选框级精修还会引入 P2/P3 投影、旋转 ROI 采样和轻量融合网络。因此，需要报告参数量、FLOPs、显存与端到端速度。精修仅作用于 post-NMS proposal，复杂度测试必须固定 max_det、置信度阈值、NMS 阈值、batch 和数值精度，并分别报告 CA 与 CA+Refine 的完整推理时间。",
    )
    replace_text(p[125], "定性结果应覆盖密集背景、遮挡、长目标、极细目标、交叉电线和边界截断。每组图同时展示真实标注、Baseline、CA 与 CA+Refine，并使用统一置信度、NMS 和 ROI 范围。失败案例也应保留，以分析漏检、断框、角度偏差及过度缩窄。")
    replace_text(p[127], "图6 不同场景下的定性检测对比。每行使用相同原图与 ROI，依次展示 Ground Truth、Baseline、CA 和 CA+Refine；对比采用统一置信度与 NMS 设置，并保留代表性失败案例。")

    replace_text(p[134], "6.3 候选框级 Refine 的作用边界与增益解释")
    replace_text(
        p[135],
        "候选框级 Refine 将精修对象从密集网格预测改为实际 post-NMS proposal，并利用方向对齐的局部特征学习实例相关尺度残差。CA 参数冻结、coarse/identity 完全恒等、残差无边界堆积以及匹配 IoU 的正向变化共同构成因果证据链。其收益主要集中于 mAP50-95 与 AP75，AP90 也获得小幅改善，而 AP95 略有下降，说明当前尺度校正仍不能保证极严格阈值下的每个样本都受益。",
    )
    replace_text(
        p[137],
        "本文仍存在三方面局限。第一，Coverage-Aware 的数值可达性不等价于特征语义充分性，尚未显式建模感受野与上下文质量。第二，精修器仅调整短、长边尺度，保持中心与角度不变，因此无法修复全部高 IoU 误差；AP95 的轻微下降也说明少量本已较准的候选框可能被过度校正。第三，当前正式结果来自单随机种子，Baseline/CA 的统一诊断、复杂度和定性结果仍需按已冻结协议补齐。上述事项属于证据完整性工作，不再触发方法结构的持续改写。",
    )
    replace_text(
        p[139],
        "本文面向低空电线等细长障碍物的旋转框检测，研究正样本分配、分布式边界回归范围与候选框局部尺度误差之间的结构性不匹配。本文提出几何可达正样本准则并据此设计 Coverage-Aware Assignment，将 DFL 表达边界前移到正样本筛选阶段；同时使用 reg_max=32 提供最小充分的长距离表达能力。在此基础上，以 CA 的 post-NMS coarse OBB 为候选框，从 P2/P3 提取方向对齐的旋转 ROI 特征，仅预测短边和长边的连续尺度残差。",
    )
    replace_text(
        p[140],
        "在固定 FP32、batch=8、imgsz=640 的独立验证中，候选框级精修将 mAP50-95 从 0.4541 提升至 0.5625，AP75 从 0.4398 提升至 0.5588，AP90 从 0.2457 提升至 0.2614；AP95 轻微下降 0.0018。恒等对照、CA 权重哈希、残差分布和匹配 IoU 分析均支持该改善来自实际几何校正。现有证据足以形成方法与实验闭环，但结论限定于当前固定验证设置，后续仅补齐统一基线、机制统计、复杂度、定性结果和必要的多种子复验。",
    )

    # Table 4: fill facts already fixed by the current protocol.
    table = doc.tables[4]
    set_cell(table.cell(3, 1), "13559 / 1695 / 待补")
    set_cell(table.cell(5, 1), "640×640")

    # Table 5: remove the obsolete CA-continue column and record only verified settings.
    table = doc.tables[5]
    remove_column(table, 3)
    headers = ["配置项", "Baseline", "CA", "CA + Refine"]
    for ci, value in enumerate(headers):
        set_cell(table.cell(0, ci), value)
    values = {
        1: ["模型规模", "YOLO11l-OBB", "YOLO11l-OBB", "CA + proposal refiner"],
        2: ["初始化权重", "待核对", "待核对", "CA best.pt"],
        3: ["训练轮数", "待核对", "300", "15（仅 Refine）"],
        4: ["输入尺寸", "640", "640", "640"],
        5: ["Batch size", "待核对", "待核对", "8"],
        6: ["优化器/初始学习率", "待核对", "待核对", "AdamW / 3×10^-4"],
        7: ["随机种子", "待核对", "待核对", "0"],
    }
    for ri, row_values in values.items():
        for ci, value in enumerate(row_values):
            set_cell(table.cell(ri, ci), value)

    # Table 6: retain one row per method and fill the audited CA/Refine results.
    table = doc.tables[6]
    remove_row(table, 4)
    remove_row(table, 3)
    set_cell(table.cell(0, 0), "方法")
    set_cell(table.cell(1, 0), "Baseline YOLO11l-OBB")
    for ci in range(1, 7):
        set_cell(table.cell(1, ci), "待补")
    set_cell(table.cell(2, 0), "CA（reg_max=32 + Coverage-Aware）")
    set_cell(table.cell(3, 0), "CA + Refine")
    for ri, variant in ((2, coarse), (3, refined)):
        for ci, key in enumerate(("precision", "recall", "map50", "map50_95", "ap75", "ap90"), start=1):
            set_cell(table.cell(ri, ci), fmt(variant[key]))
    replace_text(p[98], "表6 验证集主实验结果（FP32，batch=8）")

    # Table 9: same-checkpoint identity A/B.
    table = doc.tables[9]
    set_cell(table.cell(1, 0), "CA（coarse）")
    set_cell(table.cell(2, 0), "CA + Refine")
    set_cell(table.cell(3, 0), "差值（Refine - CA）")
    for ci, key in enumerate(("precision", "recall", "map50", "map50_95", "ap75", "ap90"), start=1):
        set_cell(table.cell(1, ci), fmt(coarse[key]))
        set_cell(table.cell(2, ci), fmt(refined[key]))
        set_cell(table.cell(3, ci), fmt(refined[key] - coarse[key], sign=True))
    replace_text(p[114], "表9 同 checkpoint 候选框级 Refine 恒等 A/B（FP32，batch=8）")

    table = doc.tables[10]
    set_cell(table.cell(3, 5), "旋转 ROI；短/长边连续残差")
    table = doc.tables[11]
    set_cell(table.cell(3, 0), "CA + Refine")
    remove_row(table, 4)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
