from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor

from myscripts.paper_submission.build_target_branches import REFERENCES, TARGETS


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "mydocs" / "创新点一" / "创新点一_覆盖能力感知与候选框级几何精修方法_数据版.docx"
DEFAULT_OUTPUT = ROOT / "mydocs" / "创新点一" / "投稿版本"


ZH_TARGETS = {
    "ivc": {
        "title": "面向无人机图像中电线旋转检测的覆盖能力感知分配与身份保持几何精修",
        "header": "面向无人机图像的电线旋转检测与几何精修",
        "abstract": (
            "低空无人机图像中的电线短轴纹理证据少、长轴跨越范围大且方向连续变化，容易暴露特征金字塔"
            "旋转检测器中的几何不相容问题：任务对齐分配可能选中语义上合适、但有限分布式回归范围无法"
            "覆盖目标全部边界的正样本。本文提出 YOLO11-OBB 的覆盖能力感知扩展。该方法在目标局部坐标系"
            "中计算候选点到边界的最大需求距离，并在任务对齐排序前过滤几何不可达候选；reg_max=32 用于"
            "为不同金字塔层提供必要的距离表达支撑。为修正剩余的高 IoU 尺度误差，本文进一步围绕 post-NMS"
            "候选框从 P2/P3 采样旋转对齐区域，仅预测短边和长边的有界对数尺度残差，同时保持中心、角度、"
            "置信度、候选数量和 NMS 身份不变。在 imgsz=640 的固定 TTPLA 验证协议下，Baseline、覆盖"
            "感知检测器和最终模型的 mAP50-95 分别为 0.4084、0.4541 和 0.5625；精修将 AP75 从 0.4398"
            "提升至 0.5588，AP90 提高 0.0157，而"
            " AP95 下降 0.0018。严格恒等对照、三种无重训练评估路径和 proposal 匹配分析表明，收益来自"
            "学习到的几何校正：匹配候选平均 IoU 提高 0.0466，63.41% 的匹配候选得到改善。该结果支持"
            "方法对电线旋转定位具有明确但有边界的正向作用。RTX 5090 上完整模型的单图耗时为"
            " 14.83±0.04 ms（67.4 FPS），相对 CA 仅增加 1.72% 参数量；正式结果来自单随机种子，"
            "因此本文不作统计显著性或机载部署声明。"
        ),
        "keywords": "电线检测；无人机图像；旋转目标检测；正样本分配；分布式回归；候选框精修",
        "note": (
            "内部投稿分支说明（投稿前删除）：本稿是 Image and Vision Computing 的后续主稿。方法冻结为"
            " Coverage-Aware Assignment + reg_max=32 表达支撑 + 身份保持的候选框级几何精修。本轮不训练"
            "模型，只使用已经审计的结果；统一 Baseline、H1/H2、复杂度、多种子、定性结果和最终测试集数据"
            "在未实测前保留为空，不用预期值替代。"
        ),
        "framing": (
            "从图像与视觉计算角度看，本文关注的核心不是无人机飞控，而是语义候选排序、有限几何表达与"
            " proposal 级定位之间的不匹配。低空无人机电线场景使这一问题更突出；本文将其作为可受控验证的"
            "旋转目标检测问题处理，不把检测结果扩展表述为闭环避障能力。"
        ),
        "limitation": (
            "面向 Image and Vision Computing，现有证据已经闭合 Baseline→CA→Refine 的主链路，但泛化"
            "范围仍然有限。后续最有价值的扩展是独立重复种子、第二个大长宽比旋转目标数据集和机载硬件"
            "测试；在这些证据完成前，不延伸为统计显著、跨域通用或部署就绪结论。"
        ),
        "conclusion": (
            "现有证据支持覆盖能力感知分配与身份保持候选框精修对电线旋转定位具有明确但有边界的正向"
            "作用，结论限定于单数据集、单正式种子和当前固定验证协议。"
        ),
    },
    "unmanned": {
        "title": "面向无人机低空感知的电线障碍物覆盖能力感知检测与候选框级几何精修",
        "header": "面向无人机低空感知的电线障碍物检测",
        "abstract": (
            "低空无人机在巡检、测绘和近基础设施飞行过程中需要可靠识别悬空电线。此类碰撞相关障碍物"
            "短边像素少、长边跨度大且方向任意，使特征金字塔旋转检测器面临正样本分配与有限分布式回归"
            "范围不匹配的问题：语义上合适的候选点未必能够完整覆盖目标四条边界。本文以 YOLO11-OBB 为"
            "基础构建电线感知模块。首先，在旋转真实框局部坐标系中计算候选点完整覆盖目标所需的最大边界"
            "距离，并在任务对齐排序前过滤几何不可达候选；同时将 reg_max 扩展为 32，为细长目标提供足够"
            "的距离表达支撑。随后，以 post-NMS 粗框为对象，从 P2/P3 提取方向对齐的局部特征，仅预测短边"
            "和长边的有界对数尺度残差，并保持中心、角度、类别置信度与候选框身份不变。在固定 imgsz=640、"
            "FP32、batch=8 的验证协议下，候选框级精修将 mAP50-95 从 0.4541 提升至 0.5625，AP75 从 0.4398"
            "提升至 0.5588；AP90 提高 0.0157，AP95 轻微下降 0.0018。恒等对照和 proposal 级分析显示，"
            "匹配框平均 IoU 增量为 0.0466，63.41% 的匹配候选得到改善。上述结果支持该方法对低空无人机"
            "电线感知具有明确但有边界的几何改进作用；本文尚未验证测距、闭环避障或机载实时部署。"
        ),
        "keywords": "无人机；低空感知；电线障碍物检测；旋转目标检测；正样本分配；几何精修",
        "note": (
            "内部投稿分支说明（投稿前删除）：目标期刊为 Unmanned Systems。正文按“无人机低空视觉感知模块”"
            "定位，投稿前须补完整 detector+NMS+Refiner 时延、参数量、FLOPs、显存和 proposal 数量相关耗时；"
            "最好增加边缘设备测试。没有闭环飞行实验时不得写成自主避障系统。"
        ),
        "framing": (
            "从无人系统角度看，本文检测器承担的是低空飞行安全链路中的视觉前端功能：将低可见度电线转换为"
            "保留方向和空间范围的旋转框输出，供后续风险评估或路径规划模块使用。现有实验只评价图像空间中的"
            "识别与几何定位，不涉及真实距离估计、飞行器动力学、航迹生成或闭环控制。"
        ),
        "limitation": (
            "面向无人机部署，当前结果仍属于离线图像感知验证。正式投稿前需测量包含解码、NMS 和旋转 ROI "
            "采样在内的完整时延；在缺少目标机载平台测试时，不应使用“实时机载”或“部署就绪”等表述。"
        ),
        "conclusion": (
            "本文输出可作为低空无人机安全感知链路中的图像空间电线 OBB，但与测距、风险判断、路径规划和"
            "飞行控制的系统集成仍需后续研究。"
        ),
    },
    "jars": {
        "title": "面向无人机低空遥感图像的电线障碍物覆盖能力感知检测与候选框级几何精修",
        "header": "面向无人机低空遥感图像的电线障碍物检测",
        "abstract": (
            "无人机低空遥感图像中的电线短边通常只有少量像素，却沿任意方向跨越较大空间范围。这种几何特征"
            "会造成特征金字塔正样本分配与分布式边界回归之间的不匹配：候选点虽然具有较高任务对齐得分，"
            "但其所在层级的有限距离表示范围可能无法覆盖完整旋转目标。本文以 YOLO11-OBB 为基础提出覆盖"
            "能力感知的电线旋转检测方法。首先，在目标局部坐标系中计算候选点覆盖四条边界所需的最大距离，"
            "并在任务对齐分配前剔除超出当前层表达范围的候选；reg_max=32 作为细长目标长距离回归的表达"
            "支撑。其次，候选框级几何精修器围绕 post-NMS 粗框采样方向对齐的 P2/P3 特征，只学习短边与长边"
            "的连续尺度残差，保持中心、角度、类别分数和候选身份不变。在固定 640 像素 FP32 验证协议下，"
            "mAP50-95 由 0.4541 提升至 0.5625，AP75 由 0.4398 提升至 0.5588；AP90 提高 0.0157，而 AP95"
            "下降 0.0018。恒等对照与 proposal 级 IoU 统计支持收益来自实例相关的尺度校正。本文为无人机"
            "低空巡检和环境测绘中的电线大长宽比旋转定位提供了一种几何建模方案。"
        ),
        "keywords": "无人机遥感；电线障碍物检测；旋转目标检测；覆盖能力感知；分布式回归；几何精修",
        "note": (
            "内部投稿分支说明（投稿前删除）：目标期刊为 Journal of Applied Remote Sensing。正文突出 UAV "
            "低空遥感、大长宽比 OBB 和空间定位价值；投稿前须补统一 Baseline/CA、H1/H2、复杂度/FPS 和"
            "全英文定性图，第二公开遥感 OBB 数据集为加分项。"
        ),
        "framing": (
            "从无人机遥感角度看，关键不仅是识别一条细线，还要恢复稳定且紧致的旋转空间范围。电线短边可能"
            "接近传感器采样极限，而长边跨越多个特征单元；其旋转框质量直接影响巡检目标映射、走廊分析、"
            "植被净空评估和后续空间测量。"
        ),
        "limitation": (
            "遥感泛化方面，当前正式证据集中于 TTPLA 低空航拍域。若要把结论扩展到更广泛的大长宽比旋转"
            "目标，应增加第二公开 OBB 数据集或跨域测试；在此之前，结论限定于当前数据和采集条件。"
        ),
        "conclusion": (
            "该方法面向无人机低空遥感图像中的电线几何定位，其对其他大长宽比目标的泛化能力仍需额外数据集验证。"
        ),
    },
    "jei": {
        "title": "面向无人机低空成像感知的电线障碍物覆盖能力感知检测与候选框级几何精修",
        "header": "面向无人机低空成像感知的电线障碍物检测",
        "abstract": (
            "低空无人机图像中的电线具有长轴跨度大、短轴纹理证据少和方向连续变化等特征，容易暴露旋转检测器"
            "中的结构性定位问题：任务对齐分配可能选择离散距离分布无法完整覆盖目标的候选点。本文提出"
            " YOLO11-OBB 的几何感知扩展。Coverage-Aware Assignment 在目标对齐坐标系中计算每个候选点"
            "所需的最大边界距离，并抑制超出特征层表达范围的候选；reg_max=32 被视为覆盖分配的表达支撑，"
            "而非独立创新。随后，候选框级精修器围绕 post-NMS 粗框提取旋转 P2/P3 区域，仅预测短边和长边"
            "的对数尺度残差。中心、角度、类别置信度和 NMS 结果均保持不变，从而可进行严格恒等对照。在"
            " imgsz=640、FP32、batch=8 的固定协议下，mAP50-95 从 0.4541 提升至 0.5625，AP75 从 0.4398"
            "提升至 0.5588。proposal 匹配结果显示平均 IoU 增量为 0.0466，63.41% 的匹配框得到改善；"
            "AP95 下降 0.0018，说明该方法获得了明确但非所有阈值一致的定位改善。"
        ),
        "keywords": "无人机成像；电线障碍物检测；旋转框；正样本分配；分布式回归；候选框精修",
        "note": (
            "内部投稿分支说明（投稿前删除）：目标期刊为 Journal of Electronic Imaging。正文突出无人机"
            "成像定位链、恒等对照、proposal 诊断和复现性；投稿前须补统一 Baseline、复杂度/完整时延和"
            "全英文可视化，单种子结果不得写成统计显著。"
        ),
        "framing": (
            "从电子成像与定位角度看，本文贡献是对旋转框定位链的受控干预：将有限距离编码、特征金字塔分配"
            "和候选框局部图像采样联系起来，并通过同权重恒等模式排除置信度重标定、二次抑制和结果回写对"
            "精度变化的干扰。"
        ),
        "limitation": (
            "面向成像系统部署，仍需在目标硬件上测量内存和完整时延。只有在 detector、解码、NMS、proposal "
            "采样与精修的整体吞吐量经过验证后，才能使用实时性表述。"
        ),
        "conclusion": (
            "该结果证明了无人机低空成像中的电线旋转框可通过受控候选级尺度校正获得改善，但完整系统速度与"
            "跨硬件复现仍需补充。"
        ),
    },
}


def _set_run_font(run, cn: str = "宋体", latin: str = "Times New Roman", size: float = 10.5, bold=None) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cn)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _replace_paragraph(paragraph: Paragraph, text: str, *, cn: str = "宋体", size: float = 10.5, bold=False) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    _set_run_font(run, cn=cn, size=size, bold=bold)


def _replace_labelled(paragraph: Paragraph, label: str, text: str) -> None:
    paragraph.clear()
    label_run = paragraph.add_run(label)
    _set_run_font(label_run, cn="黑体", size=10.5, bold=True)
    body_run = paragraph.add_run(text)
    _set_run_font(body_run, cn="宋体", size=10.5)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _insert_after(paragraph: Paragraph, text: str, *, note: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if note:
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "FFF2CC")
        p_pr.append(shd)
        run = p.add_run(text)
        _set_run_font(run, cn="黑体", size=9.5, bold=True)
        run.font.color.rgb = RGBColor(156, 87, 0)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
    else:
        try:
            p.style = "Body Text"
        except KeyError:
            pass
        run = p.add_run(text)
        _set_run_font(run, cn="宋体", size=10.5)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(21)
        p.paragraph_format.space_after = Pt(3)
    return p


def _find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError(f"paragraph not found: {prefix}")


def _clear_forced_page_break(paragraph: Paragraph) -> None:
    """Remove a paragraph-level page break inherited from the source DOCX."""
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:pageBreakBefore"))
    if node is not None:
        p_pr.remove(node)


def _set_table(table, rows: list[list[str]], *, add_columns: int = 0) -> None:
    """Replace a table while preserving the source document's layout container."""
    for _ in range(add_columns):
        table.add_column(Inches(0.72))
    if len(rows) != len(table.rows):
        raise ValueError(f"table row mismatch: expected {len(table.rows)}, received {len(rows)}")
    if any(len(row) != len(table.columns) for row in rows):
        raise ValueError(f"table column mismatch: expected {len(table.columns)}")
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _set_run_font(
                        run,
                        cn="黑体" if row_index == 0 else "宋体",
                        size=8.0 if len(table.columns) >= 7 else 8.5,
                        bold=row_index == 0,
                    )


def _replace_ivc_final_content(doc: Document) -> None:
    """Freeze the Chinese IVC branch to the audited 2026-08-16 evidence."""
    table_rows = {
        0: [
            ["长边分桶", "Baseline 正样本", "Baseline 溢出率/%", "CA 正样本", "CA 溢出率/%"],
            ["<100 px", "3038", "0.0", "3038", "0.0"],
            ["100–200 px", "3369", "3.4", "3370", "0.0"],
            ["200–300 px", "4681", "20.9", "4686", "0.0"],
            ["300–500 px", "5172", "55.3", "5176", "0.0"],
            [">500 px", "18727", "99.9", "18706", "0.2"],
        ],
        4: [
            ["项目", "数值"],
            ["处理后训练图像数", "13559"],
            ["处理后验证图像数", "1695"],
            ["验证集实例数", "3603"],
            ["输入尺寸", "640×640"],
            ["标注形式", "电线旋转框（OBB）"],
            ["主结果用途", "固定验证集统一评估"],
        ],
        5: [
            ["配置项", "Baseline", "CA", "CA + Refine"],
            ["模型/权重", "YOLO11l-OBB", "CA best.pt", "CA + epoch 4 Refiner"],
            ["输入尺寸", "640", "640", "640"],
            ["验证精度", "FP32", "FP32", "FP32"],
            ["Batch size", "8", "8", "8"],
            ["conf / NMS IoU", "0.01 / 0.70", "0.01 / 0.70", "0.01 / 0.70"],
            ["max_det", "300", "300", "300"],
            ["Refiner 训练", "—", "—", "15 epoch，seed 0"],
        ],
        6: [
            ["方法", "Precision", "Recall", "mAP50", "mAP50-95", "AP75", "AP90"],
            ["Baseline YOLO11l-OBB", "0.7901", "0.7777", "0.7254", "0.4084", "0.4370", "0.0971"],
            ["CA（reg_max=32 + Coverage-Aware）", "0.8044", "0.7738", "0.7395", "0.4541", "0.4398", "0.2457"],
            ["CA + Refine", "0.8939", "0.8665", "0.8936", "0.5625", "0.5588", "0.2614"],
        ],
        7: [
            ["诊断项", "Baseline", "CA", "变化"],
            ["全体正样本 DFL 溢出率/%", "64.8", "0.11", "−64.7 pp"],
            ["长边 >500 px 溢出率/%", "99.9", "0.2", "−99.7 pp"],
            ["长边 >500 px 的 P3 占比/%", "1.1", "0.2", "−0.9 pp"],
            ["长边 >500 px 的 P4 占比/%", "98.9", "98.9", "0.0 pp"],
            ["长边 >500 px 的 P5 占比/%", "0.1", "0.9", "+0.8 pp"],
            ["单侧理论范围 P3/P4/P5", "120/240/480", "248/496/992", "reg_max 16→32"],
        ],
        8: [
            ["候选级诊断", "数值", "判定", "说明"],
            ["有效/匹配 proposal", "5260 / 3444", "通过", "固定验证集"],
            ["平均匹配 IoU 增量", "+0.0466", "正向", "Refine − coarse"],
            ["改善/恶化比例", "63.41% / 36.59%", "正向", "改善比例更高"],
            ["短/长边边界命中", "0 / 0", "通过", "无边界塌缩"],
        ],
        9: [
            ["推理模式", "Precision", "Recall", "mAP50", "mAP50-95", "AP75", "AP90"],
            ["CA（coarse）", "0.8044", "0.7738", "0.7395", "0.4541", "0.4398", "0.2457"],
            ["Identity（零残差）", "0.8044", "0.7738", "0.7395", "0.4541", "0.4398", "0.2457"],
            ["CA + Refine", "0.8939", "0.8665", "0.8936", "0.5625", "0.5588", "0.2614"],
        ],
        10: [
            ["方法", "特征增强", "旋转框", "覆盖感知分配", "reg_max 扩容", "连续残差精修", "mAP50", "mAP50-95", "FPS"],
            ["PL-YOLOv8[1]", "方向滤波块", "是", "否", "未显式讨论", "否", "跨协议", "不直接比较", "跨硬件"],
            ["本文 CA 主干", "否", "是", "是", "32", "否", "0.7395", "0.4541", "80.65"],
            ["本文最终模型", "否", "是", "是", "32", "短/长边连续残差", "0.8936", "0.5625", "67.42"],
        ],
        11: [
            ["方法", "参数量/M", "FLOPs/G", "显存/GB", "延迟/ms", "FPS"],
            ["Baseline", "26.160", "90.970", "0.259", "12.44±0.28", "80.37"],
            ["CA", "27.267", "96.111", "0.263", "12.40±0.18", "80.65"],
            ["CA + Refine", "27.736", "≥96.761", "0.304", "14.83±0.04", "67.42"],
        ],
    }
    for index, rows in table_rows.items():
        _set_table(doc.tables[index], rows)
    _set_table(
        doc.tables[1],
        [
            ["长边分桶", "Base P3/%", "Base P4/%", "Base P5/%", "CA P3/%", "CA P4/%", "CA P5/%"],
            ["100–200 px", "81.7", "18.3", "0.0", "99.9", "0.1", "0.1"],
            ["200–300 px", "26.4", "73.6", "0.0", "99.1", "0.9", "0.0"],
            ["300–500 px", "4.0", "96.0", "0.0", "66.5", "33.5", "0.0"],
            [">500 px", "1.1", "98.9", "0.1", "0.2", "98.9", "0.9"],
        ],
        add_columns=3,
    )

    replacements = {
        "第二类统计正样本层级分布": "第二类统计正样本层级分布，分别计算不同长度桶中正样本落在 P3、P4、P5 的比例。该统计用于解释溢出的层级来源，但不预设长目标必须迁移至 P5：若某一层级的分配占比不变而扩展 reg_max 后溢出消失，则说明关键矛盾是该层有限表示范围，而非简单的层级误分配。",
        "Coverage-Aware Assignment 主要修复": "Coverage-Aware Assignment 主要修复‘长目标能否在当前层级的离散范围内被完整表达’，但完整覆盖不等于高精度贴合。对于极细旋转框，短边误差、法向中心偏移和角度误差都会导致 IoU 快速下降。本文最终通过 coarse/identity/refined 恒等对照、匹配候选 IoU 变化和残差边界统计区分精修收益，避免把分数重标定或固定缩放误认为几何学习。",
        "实验采用 TTPLA 系列低空航拍电力线数据": "实验采用 TTPLA 低空航拍电力线数据[5]。高分辨率图像以 640×640 切片表示，并保留电线旋转几何标注。固定处理后划分包含 13559 张训练图像和 1695 张验证图像，验证集共有 3603 个标注实例。所有方法共享相同图像、标注转换和预处理；Refiner 的 checkpoint 仅在训练集内部的确定性 image-level holdout 上选择，论文数值统一在未参与选择的验证集上重新计算。",
        "基线采用 YOLO11l-OBB": "基线采用 YOLO11l-OBB，CA 模型使用 Coverage-Aware Assignment 并设 reg_max=32，最终模型在冻结 CA 后连接候选框级几何精修器。Refiner 使用 P2/P3 特征、32 个 ROI 通道、5×24 旋转 ROI 和 128 维隐藏层，训练 15 个 epoch；优化器为 AdamW，初始学习率 3×10^-4，weight decay 为 1×10^-4，warmup 为 3 个 epoch，随机种子为 0。正式验证统一采用 imgsz=640、FP32、batch=8、conf=0.01、NMS IoU=0.70 和 max_det=300。",
        "主结果以相同验证设置比较": "主结果在完全一致的固定验证设置下比较 Baseline、CA 与 CA+Refine。Baseline 和 CA 使用各自冻结权重，最终模型在冻结 CA 后仅加载所选 Refiner；coarse 与 identity 只用于恒等性校验，不作为独立方法重复计入主表。",
        "Coverage-Aware 的直接证据": "Coverage-Aware 的直接证据来自正样本可达性与层级分布。Baseline 的总体 DFL 溢出率约为 64.8%，CA 降至 0.11%；长边超过 500 px 时由 99.9% 降至 0.2%。层级统计同时表明，Baseline 的极长目标本就主要位于 P4，问题并非简单的‘长目标错误进入 P3’，而是 reg_max=16 下 P4 的有限距离范围仍不足。CA 基本消除该溢出，并使部分中长目标在范围可达时保留于更高分辨率的 P3。",
        "除聚合 AP 外": "除聚合 AP 外，本文从候选框匹配层面检查精修是否真正改善几何。固定验证中共有 5260 个有效 proposal，其中 3444 个与真实框匹配；匹配 proposal 的平均 IoU 增量为 0.0466，改善比例为 63.41%，高于 36.59% 的恶化比例。短边与长边残差边界命中率均为 0；表8汇总这些候选级诊断。",
        "reg_max 扩容会增加": "reg_max 扩容会增加边界分布输出通道，候选框级精修还会引入 P2/P3 投影、旋转 ROI 采样和轻量融合网络。在 RTX 5090、FP32、batch=1 的完整链路测试中，CA+Refine 参数量为 27.736 M，单图延迟 14.83±0.04 ms，吞吐率为 67.42 FPS；相对 CA 增加 0.469 M 参数和 2.43 ms。Refiner FLOPs 随 proposal 数变化，表中 ≥96.761 G 为两个 proposal 下的实测下界。Baseline 的跨轮波动为 5.43%，因此 Baseline 与 CA 的细小时延差异仅作描述，不作速度优劣结论。",
        "定性结果应覆盖": "定性图由冻结验证导出自动选择，所有列使用相同图像、阈值与 NMS 设置。图6保留两组明显改善、一组近中性以及一组负向案例，用于同时展示局部尺度校正的有效情形与边界。绿色为 GT，红色为预测。",
        "本文仍存在三方面局限": "本文仍存在四方面局限。第一，数值可达性不等价于特征语义充分性。第二，精修器只调整短、长边，不能修复中心和角度误差，AP95 的轻微下降也表明少量高精度候选可能被过度校正。第三，正式结果来自单随机种子，三条复现路径复用同一 checkpoint，不能替代独立重复训练。第四，现有证据集中于一个 TTPLA 处理设置和一张桌面 GPU，尚不能支持跨域泛化、统计显著性或机载部署结论。",
        "在固定 FP32、batch=8、imgsz=640": "在固定 FP32、batch=8、imgsz=640 的独立验证中，Baseline、CA 与 CA+Refine 的 mAP50-95 分别为 0.4084、0.4541 和 0.5625；精修相对 CA 将 AP75 从 0.4398 提升至 0.5588，AP90 从 0.2457 提升至 0.2614，同时 AP95 轻微下降 0.0018。恒等对照、CA 权重哈希、残差边界统计和匹配 IoU 分析均支持主要改善来自实际几何校正。现有证据已形成 Baseline→CA→Refine 的方法与实验闭环，但结论限定于当前固定验证和单随机种子设置。",
    }
    for prefix, text in replacements.items():
        _replace_paragraph(_find_paragraph(doc, prefix), text, cn="宋体", size=10.5)
    _replace_paragraph(_find_paragraph(doc, "表1 原始分配"), "表1 Baseline 与 CA 的 DFL 覆盖诊断", cn="宋体", size=9)
    _replace_paragraph(_find_paragraph(doc, "表2 原始分配"), "表2 Baseline 与 CA 的正样本层级分布", cn="宋体", size=9)
    _replace_paragraph(_find_paragraph(doc, "表8 Oracle"), "表8 候选框级几何精修诊断", cn="宋体", size=9)


def build(target_key: str, output_dir: Path, *, final: bool = False) -> Path:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    meta = ZH_TARGETS[target_key]
    en_meta = TARGETS[target_key]
    output_dir.mkdir(parents=True, exist_ok=True)
    final_mode = final and target_key == "ivc"
    suffix = "投稿完整稿" if final_mode else "投稿预备分支"
    out = output_dir / f"{en_meta['file_stem']}_{suffix}_中文版.docx"
    shutil.copy2(SOURCE, out)

    doc = Document(out)
    nonempty = [p for p in doc.paragraphs if p.text.strip()]
    _replace_paragraph(nonempty[0], meta["title"], cn="黑体", size=18, bold=True)
    nonempty[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Keep the running header consistent with the target-specific body title.
    # Some source templates share the same header part across sections, so track
    # part identities to avoid replacing the same paragraph repeatedly.
    seen_header_parts: set[int] = set()
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            header_part_id = id(header.part)
            if header_part_id in seen_header_parts:
                continue
            seen_header_parts.add(header_part_id)
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    _replace_paragraph(paragraph, meta["header"], cn="宋体", size=8.5)
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    author = _find_paragraph(doc, "作者：")
    if not final_mode:
        _insert_after(author, meta["note"], note=True)

    abstract = _find_paragraph(doc, "摘要：")
    _replace_labelled(abstract, "摘要：", meta["abstract"])
    keywords = _find_paragraph(doc, "关键词：")
    _replace_labelled(keywords, "关键词：", meta["keywords"])

    english_abstract = _find_paragraph(doc, "Power lines")
    _replace_paragraph(english_abstract, en_meta["abstract"], cn="Times New Roman", size=9.5)
    english_abstract.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    english_keywords = _find_paragraph(doc, "Key words:")
    _replace_paragraph(english_keywords, f"Key words: {en_meta['keywords']}", cn="Times New Roman", size=9.5)

    intro_first = _find_paragraph(doc, "电线、拉线和细杆等低空障碍物")
    framing_anchor = _insert_after(intro_first, meta["framing"])
    if target_key == "ivc":
        uav_anchor = _insert_after(
            framing_anchor,
            "低空无人机目标检测综述已归纳尺度、视角、遮挡与数据差异等问题[14]；已有无人机视觉基准进一步表明，小目标、相机运动和视角变化会持续削弱检测稳定性[22]。对于短边仅由少量像素定义的目标，IoU 对轻微坐标偏移也格外敏感[23]。这些一般性困难在电线场景中叠加出现，使多尺度分配与高 IoU 定位不能只依赖更强的全局语义特征。",
        )
        ivc_anchor = _insert_after(
            uav_anchor,
            "与近期 Image and Vision Computing 论文相比，Bai 等侧重自注意引导、全局特征融合与小目标分配[15]，Sang 等侧重环境自适应上下文和快速检测[16]，Chaurasia 与 Patro 侧重通道—空间注意及旋转角分类[17]；Rong 等则从方向特征增强切入电线检测[1]。本文不直接横向比较跨数据集 AP，而是把差异限定为算法问题：正样本是否在有限距离分布下几何可表示，以及 post-NMS 旋转候选能否在保持身份与置信度不变的前提下获得局部尺度校正。",
        )
        assignment_anchor = _insert_after(
            ivc_anchor,
            "FPN 为多尺度检测提供基础层级表示[18]，Mask R-CNN 说明 proposal 特征能够支持第二次几何决策[19]，DOTA 则系统暴露了航拍目标的尺度、方向和形状变化[20]。在正样本分配方面，ATSS 与 OTA 分别从自适应统计和全局最优传输角度说明正负样本定义会直接影响检测器学习[24,25]；在旋转目标定位方面，CSL、KLD、GWD、RoI Transformer、ReDet、Oriented R-CNN 与 R3Det 分别处理角度周期性、旋转框度量、方向对齐 proposal 或迭代精修[21,26-31]。针对大长宽比旋转框的最新研究也表明，该问题仍未完全解决[13]。本文的区别是把有限 DFL 距离支撑显式引入候选可达性判断，并将最终精修严格限制为 post-NMS 候选的短边和长边尺度更新。",
        )
        _insert_after(
            assignment_anchor,
            "近期电线视觉方法覆盖实例分割、多任务检测、轻量组件检测和形状感知分割等路线：CableNet 保留电缆实例身份[32]，PowerLine-MTYOLO 联合电缆分割与断股检测[33]，LPC-Det 面向无人机电力线组件的轻量检测[34]，SFFPLDN 融合形状感知与多尺度特征[35]。这些工作进一步说明特征、任务和效率设计的重要性，但并未直接回答本文所关注的层级回归可达性与同身份候选尺度校正问题。",
        )

    limitation = _find_paragraph(doc, "本文仍存在三方面局限")
    _insert_after(limitation, meta["limitation"])

    conclusion_last = _find_paragraph(doc, "在固定 FP32、batch=8、imgsz=640")
    _insert_after(conclusion_last, meta["conclusion"])

    if target_key == "ivc":
        identity_result = _find_paragraph(doc, "表9表明，coarse 与 identity")
        _insert_after(
            identity_result,
            "无重训练评估路径复核如下：FP32、batch=8 时 coarse/refined 为 0.454137/0.562504，增量 0.108368；FP32、batch=1 时为 0.454151/0.562460，增量 0.108310；AMP、batch=8 时为 0.453562/0.562017，增量 0.108455。三条路径的增量极差仅 0.000146，说明主提升不依赖单一 batch 或数值精度设置。该对照复用同一 checkpoint，只属于评估链稳健性测试，不能视为独立种子或统计显著性证据。",
        )

        # Normalize the existing source references and append the new IVC
        # bibliography entries. The English citation text is retained so the
        # Chinese and English review branches share one auditable source list.
        reference_paragraphs = {
            int(match.group(1)): paragraph
            for paragraph in doc.paragraphs
            if (match := re.match(r"^\[(\d+)\]", paragraph.text.strip()))
        }
        for index, reference in enumerate(REFERENCES, start=1):
            text = f"[{index}] {reference}"
            paragraph = reference_paragraphs.get(index)
            if paragraph is None:
                paragraph = doc.add_paragraph()
            _replace_paragraph(paragraph, text, cn="Times New Roman", size=8)
            paragraph.paragraph_format.left_indent = Pt(12)
            paragraph.paragraph_format.first_line_indent = Pt(-12)
            paragraph.paragraph_format.space_after = Pt(1.5)

        # The source data draft deliberately separated incomplete tables onto
        # new pages. In the IVC review branch this created a nearly blank page
        # before Table 11, so keep the caption with its table but remove the
        # inherited forced page break.
        _clear_forced_page_break(_find_paragraph(doc, "表11 模型复杂度"))
        if final_mode:
            _replace_ivc_final_content(doc)
            # The source placeholders for Figs. 3, 5, and 6 used a shallow
            # landscape ratio.  Match the final exported assets so Word does
            # not vertically compress the charts or qualitative grid.
            final_figure_sizes = {
                2: (6.10, 2.70),
                4: (6.10, 2.68),
                5: (6.10, 6.22),
            }
            for index, (width, height) in final_figure_sizes.items():
                doc.inline_shapes[index].width = Inches(width)
                doc.inline_shapes[index].height = Inches(height)

    props = doc.core_properties
    props.title = meta["title"]
    props.subject = f"{en_meta['journal']} 中文完整投稿稿" if final_mode else f"{en_meta['journal']} 中文投稿审阅分支"
    props.keywords = meta["keywords"]
    props.comments = "实验数据已补齐；作者、单位和披露信息仍需提交前填写。" if final_mode else "内部中文审阅分支；投稿前删除黄色说明并按目标期刊模板处理。"
    doc.save(out)
    # Remove stale font-table declarations that can make Word report a
    # non-used Japanese fallback (e.g. MS Gothic) in the font inspector.
    from lxml import etree
    from zipfile import ZIP_DEFLATED, ZipFile

    temp = out.with_suffix(".fontfix.docx")
    with ZipFile(out, "r") as src, ZipFile(temp, "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename == "word/fontTable.xml":
                root = OxmlElement("w:fonts")
                for name in ("Times New Roman", "宋体", "黑体", "Cambria Math"):
                    font = OxmlElement("w:font")
                    font.set(qn("w:name"), name)
                    root.append(font)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            dst.writestr(item, data)
    temp.replace(out)
    return out


def main() -> None:
    parser = argparse.ArgumentParser(description="Build Chinese review counterparts for target-specific journal branches.")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--target", choices=[*ZH_TARGETS, "all"], default="all")
    parser.add_argument("--final", action="store_true", help="生成证据已补齐的 IVC 中文完整稿")
    args = parser.parse_args()
    keys = ZH_TARGETS if args.target == "all" else [args.target]
    for key in keys:
        print(build(key, args.output_dir, final=args.final))


if __name__ == "__main__":
    main()
