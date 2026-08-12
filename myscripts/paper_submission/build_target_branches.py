from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE_FIGURES = ROOT / "mydocs" / "创新点一" / "paper_visuals" / "outputs"
IVC_FIGURES = ROOT / "mydocs" / "创新点一" / "投稿版本" / "IVC_assets"
ACTIVE_IVC = False
BODY_FONT_SIZE = 9.5


CA = {
    "Precision": 0.804378,
    "Recall": 0.773765,
    "mAP@0.50": 0.739488,
    "mAP@0.50:0.95": 0.454137,
    "AP@0.75": 0.439822,
    "AP@0.90": 0.245697,
    "AP@0.95": 0.099024,
}
REFINE = {
    "Precision": 0.893915,
    "Recall": 0.866500,
    "mAP@0.50": 0.893567,
    "mAP@0.50:0.95": 0.562504,
    "AP@0.75": 0.558800,
    "AP@0.90": 0.261374,
    "AP@0.95": 0.097174,
}


TARGETS = {
    "ivc": {
        "journal": "Image and Vision Computing",
        "short": "IVC",
        "file_stem": "IVC",
        "title": (
            "Coverage-Aware Assignment and Identity-Preserving Geometric Refinement for "
            "Oriented Power-Line Detection in UAV Imagery"
        ),
        "abstract": (
            "Power lines are difficult targets in low-altitude unmanned aerial vehicle imagery because their short "
            "axis contains little appearance evidence, their long axis can span a large image region, and their "
            "orientation varies continuously. These properties expose a geometric incompatibility in feature-pyramid "
            "oriented detectors: task-aligned assignment can select a semantically suitable positive whose finite "
            "distributional regression support cannot reach all target boundaries. We introduce a coverage-aware "
            "extension of YOLO11-OBB that evaluates the candidate-to-boundary demand in the target-local coordinate "
            "frame and filters geometrically infeasible positives before task-aligned ranking. A distributional "
            "support of reg_max=32 is used to make feasible assignments available across pyramid levels. To correct "
            "the remaining high-IoU scale error, an identity-preserving proposal refiner samples rotated P2/P3 regions "
            "around post-NMS boxes and predicts bounded short- and long-side log-scale residuals while preserving "
            "center, angle, confidence, proposal count, and NMS identity. On the fixed TTPLA validation protocol with "
            "640-pixel inputs, refinement improves mAP@0.50:0.95 from 0.4541 to 0.5625 and AP@0.75 from 0.4398 to "
            "0.5588. AP@0.90 increases by 0.0157, whereas AP@0.95 decreases by 0.0018. Exact identity controls, three "
            "evaluation-path checks, and proposal matching attribute the gain to learned geometric correction: mean "
            "matched-proposal IoU increases by 0.0466 and 63.41% of matched proposals improve. The results support a "
            "controlled improvement in oriented power-line localization, while repeated-seed, complexity, and final "
            "test-set evidence remain to be completed."
        ),
        "keywords": (
            "power-line detection; UAV imagery; oriented object detection; positive-sample assignment; "
            "distributional regression; proposal refinement"
        ),
        "framing": (
            "For an image-and-vision audience, the core problem is a mismatch between semantic candidate ranking, "
            "finite geometric representation, and proposal-specific localization. The UAV power-line setting is the "
            "application domain that makes this mismatch observable; the contribution is evaluated as a controlled "
            "computer-vision intervention rather than as an end-to-end flight-control system."
        ),
        "target_note": (
            "This is the primary manuscript branch for Image and Vision Computing. The method is frozen as "
            "coverage-aware assignment + reg_max=32 support + identity-preserving proposal-level refinement. No "
            "model training is authorized in this drafting cycle. Use only existing audited results; leave Baseline, "
            "mechanism, complexity, multi-seed, qualitative, and final test evidence visibly blank until measured."
        ),
    },
    "unmanned": {
        "journal": "Unmanned Systems",
        "short": "US",
        "file_stem": "Unmanned_Systems",
        "title": (
            "Coverage-Aware Oriented Detection and Proposal-Level Geometric Refinement of "
            "Power-Line Obstacles for Low-Altitude UAV Perception"
        ),
        "abstract": (
            "Reliable perception of suspended power lines is important for low-altitude unmanned aerial vehicle "
            "operations because these collision-relevant obstacles are visually thin, extend across a large image "
            "region, and appear at arbitrary orientations. Their geometry exposes a mismatch in feature-pyramid "
            "oriented detectors: task-aligned assignment may select a semantically suitable candidate whose finite "
            "distributional regression range cannot cover all boundaries of the target. This paper develops a "
            "power-line perception module based on YOLO11-OBB. Coverage-aware assignment evaluates the boundary "
            "distance required by each candidate in the local coordinate frame of the rotated target and removes "
            "geometrically infeasible candidates before task-aligned ranking. The distributional support is expanded "
            "to reg_max=32. A proposal-level geometric refiner then samples orientation-aligned P2/P3 features around "
            "post-NMS coarse boxes and predicts bounded short- and long-side log-scale residuals while preserving "
            "center, angle, class confidence, and proposal identity. Under a fixed 640-pixel FP32 validation protocol, "
            "the refiner increases mAP@0.50:0.95 from 0.4541 to 0.5625 and AP@0.75 from 0.4398 to 0.5588. AP@0.90 "
            "increases by 0.0157, whereas AP@0.95 decreases by 0.0018. Identity controls and proposal-level analysis "
            "show a mean matched-IoU change of +0.0466, with 63.41% of matched proposals improved. The results support "
            "the method as a bounded geometric improvement to a UAV perception pipeline; they do not constitute a "
            "closed-loop collision-avoidance or onboard real-time demonstration."
        ),
        "keywords": (
            "unmanned aerial vehicle; low-altitude perception; power-line obstacle detection; oriented object "
            "detection; positive-sample assignment; geometric refinement"
        ),
        "framing": (
            "For an unmanned-systems audience, the detector is positioned as the visual perception component that "
            "converts low-observability wire evidence into an oriented geometric representation. The oriented box "
            "retains direction and spatial extent that can be consumed by a later navigation or risk-assessment "
            "module, but the present study evaluates perception accuracy only and does not claim closed-loop flight "
            "control, metric ranging, or autonomous avoidance."
        ),
        "target_note": (
            "This is the primary UAV-perception branch. Before submission, add complete detector-plus-refiner latency, "
            "parameters, FLOPs, peak memory, and proposal-count-dependent timing. An onboard or edge-device benchmark "
            "is strongly preferred. Keep all claims at perception-module level unless closed-loop flight experiments "
            "are added."
        ),
    },
    "jars": {
        "journal": "Journal of Applied Remote Sensing",
        "short": "JARS",
        "file_stem": "JARS",
        "title": (
            "Coverage-Aware Assignment and Proposal-Level Geometric Refinement for "
            "Power-Line Obstacle Detection in Low-Altitude UAV Imagery"
        ),
        "abstract": (
            "Power lines are visually inconspicuous obstacles in low-altitude unmanned aerial vehicle imagery. They "
            "occupy only a few pixels across their short axis while "
            "extending over a large spatial range and arbitrary orientations. These characteristics create a "
            "geometric mismatch between feature-pyramid assignment and distributional box regression: a candidate "
            "may be semantically suitable yet unable to represent all four distances to an oriented target within "
            "the finite regression range of its feature level. This paper introduces a coverage-aware oriented "
            "detector built on YOLO11-OBB. First, the distance required to cover each ground-truth box is computed in "
            "the local coordinate system of the rotated target, and candidates outside the representable range are "
            "removed before task-aligned assignment. The distributional range is expanded to reg_max=32 to provide "
            "sufficient support for elongated targets. Second, a proposal-level geometric refiner samples "
            "orientation-aligned P2/P3 features around post-NMS coarse boxes and predicts only short- and long-side "
            "log-scale residuals, preserving center, angle, class confidence, and proposal identity. Under a fixed "
            "640-pixel FP32 validation protocol, refinement improves mAP@0.50:0.95 from 0.4541 to 0.5625 and AP@0.75 "
            "from 0.4398 to 0.5588; AP@0.90 increases from 0.2457 to 0.2614, whereas AP@0.95 decreases slightly by "
            "0.0018. Identity controls and proposal-level IoU analysis attribute the gain to instance-dependent scale "
            "correction. The method provides a geometry-centered power-line perception component for low-altitude "
            "UAV inspection and environmental mapping."
        ),
        "keywords": "UAV remote sensing; power-line obstacle detection; oriented object detection; assignment; distribution focal loss; geometric refinement",
        "framing": (
            "For a remote-sensing audience, the central problem is not merely detecting a thin visual pattern in a "
            "UAV frame. It is "
            "recovering a stable oriented footprint for an elongated asset whose projected width may approach the "
            "sampling limit of the sensor while its length crosses several feature-map cells. Such geometry affects "
            "mapping, corridor inspection, vegetation-clearance assessment, and downstream spatial measurements."
        ),
        "target_note": (
            "This branch emphasizes low-altitude aerial remote sensing, large-aspect-ratio oriented targets, and "
            "spatially meaningful localization. Before submission, add end-to-end complexity, an English qualitative "
            "figure, and the frozen Baseline/CA mechanism statistics. A second public remote-sensing OBB dataset is "
            "desirable but is not fabricated in this draft."
        ),
    },
    "jei": {
        "journal": "Journal of Electronic Imaging",
        "short": "JEI",
        "file_stem": "JEI",
        "title": (
            "Coverage-Aware Assignment and Candidate-Level Geometric Refinement for "
            "Power-Line Obstacle Detection in Low-Altitude UAV Perception"
        ),
        "abstract": (
            "Power-line obstacles are difficult to localize in low-altitude unmanned aerial vehicle images because their "
            "long axis spans many pixels, their short axis contains little appearance evidence, and their orientation "
            "varies continuously. In a feature-pyramid oriented detector, these properties expose a structural "
            "mismatch: task-aligned assignment can select a candidate whose discrete distance distribution cannot "
            "cover the complete target. We present a geometry-aware extension of YOLO11-OBB. Coverage-aware assignment "
            "computes the maximum boundary distance required by each candidate in the target-aligned coordinate frame "
            "and suppresses candidates that exceed the representable range of their feature level. The regression "
            "distribution is enlarged to reg_max=32 as an implementation support rather than an independent novelty. "
            "A proposal-level refiner then extracts rotated P2/P3 regions around post-NMS coarse boxes and predicts "
            "only short- and long-side log-scale residuals. The center, angle, class confidence, and NMS result remain "
            "unchanged, enabling an exact identity control. With 640-pixel inputs and a fixed FP32 batch-eight protocol, "
            "mAP@0.50:0.95 increases from 0.4541 to 0.5625 and AP@0.75 from 0.4398 to 0.5588. Proposal matching shows a "
            "mean IoU increment of 0.0466, with 63.41% of matched proposals improved. AP@0.95 decreases by 0.0018, so "
            "the evidence supports a clear but bounded localization improvement rather than uniform gains at every "
            "IoU threshold."
        ),
        "keywords": "UAV perception; power-line obstacle detection; oriented bounding box; positive-sample assignment; distributional regression; proposal refinement",
        "framing": (
            "For an electronic-imaging audience, the contribution is a controlled intervention in the localization "
            "pipeline. The method links finite-distance encoding, feature-pyramid assignment, and candidate-level "
            "image sampling, while identity controls isolate the effect of geometric refinement from confidence "
            "recalibration and secondary suppression."
        ),
        "target_note": (
            "This branch emphasizes imaging-system localization, diagnostic controls, and reproducibility. Before "
            "submission, add the unified baseline, parameter/FLOP/latency measurements, and English visual results. "
            "The present single-seed result must remain described as a controlled result, not a significance claim."
        ),
    },
}


REFERENCES = [
    "S. Rong, L. He, S. F. Atici, and A. E. Cetin, ‘Advanced YOLO-based real-time power line detection for vegetation management,’ IEEE Transactions on Power Delivery 40(4), 2142–2153 (2025). https://doi.org/10.1109/TPWRD.2025.3578274.",
    "L. Yang, J. Fan, Y. Liu, E. Li, J. Peng, and Z. Liang, ‘A review on state-of-the-art power line inspection techniques,’ IEEE Transactions on Instrumentation and Measurement 69(12), 9350–9365 (2020). https://doi.org/10.1109/TIM.2020.3031194.",
    "Ö. E. Yetgin, B. Benligiray, and Ö. N. Gerek, ‘Power line recognition from aerial images with deep learning,’ IEEE Transactions on Aerospace and Electronic Systems 55(5), 2241–2252 (2019). https://doi.org/10.1109/TAES.2018.2883879.",
    "P. Sharma, S. Saurav, and S. Singh, ‘Object detection in power line infrastructure: a review of the challenges and solutions,’ Engineering Applications of Artificial Intelligence 130, 107781 (2024). https://doi.org/10.1016/j.engappai.2023.107781.",
    "R. Abdelfattah, X. Wang, and S. Wang, ‘TTPLA: an aerial-image dataset for detection and segmentation of transmission towers and power lines,’ in Proceedings of the Asian Conference on Computer Vision (2020). https://doi.org/10.1007/978-3-030-63486-5_29.",
    "J. Redmon, S. Divvala, R. Girshick, and A. Farhadi, ‘You only look once: unified, real-time object detection,’ in Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition, 779–788 (2016). https://doi.org/10.1109/CVPR.2016.91.",
    "Ultralytics, ‘Ultralytics YOLO11,’ https://docs.ultralytics.com/models/yolo11 (accessed 12 July 2026).",
    "C. Feng, Y. Zhong, Y. Gao, M. R. Scott, and W. Huang, ‘TOOD: task-aligned one-stage object detection,’ in Proceedings of the IEEE/CVF International Conference on Computer Vision, 3510–3519 (2021).",
    "X. Li, W. Wang, L. Wu, S. Chen, X. Hu, J. Li, J. Tang, and J. Yang, ‘Generalized focal loss: learning qualified and distributed bounding boxes,’ Advances in Neural Information Processing Systems 33, 21002–21012 (2020).",
    "J. Murrugarra-Llerena, L. N. Kirsten, L. F. Zeni, and C. R. Jung, ‘Probabilistic intersection-over-union for training and evaluation of oriented object detectors,’ IEEE Transactions on Image Processing 33, 671–681 (2024). https://doi.org/10.1109/TIP.2023.3348697.",
    "M. He, L. Qin, X. Deng, S. Zhou, H. Liu, and K. Liu, ‘Transmission line segmentation solutions for UAV aerial photography based on improved U-Net,’ Drones 7(4), 274 (2023). https://doi.org/10.3390/drones7040274.",
    "D. Bolya, C. Zhou, F. Xiao, and Y. J. Lee, ‘YOLACT++: better real-time instance segmentation,’ IEEE Transactions on Pattern Analysis and Machine Intelligence 44(2), 1108–1121 (2022). https://doi.org/10.1109/TPAMI.2020.3014297.",
    "Y. Tan, L. Deng, and D. Zhao, ‘Unified method for oriented object detection with large aspect ratio and square-like object,’ Journal of Applied Remote Sensing 20(2), 021407 (2026). https://doi.org/10.1117/1.JRS.20.021407.",
    "P. Mittal, R. Singh, and A. Sharma, ‘Deep learning-based object detection in low-altitude UAV datasets: a survey,’ Image and Vision Computing 104, 104046 (2020). https://doi.org/10.1016/j.imavis.2020.104046.",
    "J. Bai, H. Hu, X. Liu, S. Zhuang, and Z. Wang, ‘UAV image object detection based on self-attention guidance and global feature fusion,’ Image and Vision Computing 151, 105262 (2024). https://doi.org/10.1016/j.imavis.2024.105262.",
    "M. Sang, S. Tian, L. Yu, G. Wang, and Y. Peng, ‘Environmentally adaptive fast object detection in UAV images,’ Image and Vision Computing 148, 105103 (2024). https://doi.org/10.1016/j.imavis.2024.105103.",
    "D. Chaurasia and B. D. K. Patro, ‘Detection of objects in satellite and aerial imagery using channel and spatially attentive YOLO-CSL for surveillance,’ Image and Vision Computing 147, 105070 (2024). https://doi.org/10.1016/j.imavis.2024.105070.",
    "T.-Y. Lin, P. Dollár, R. Girshick, K. He, B. Hariharan, and S. Belongie, ‘Feature pyramid networks for object detection,’ in 2017 IEEE Conference on Computer Vision and Pattern Recognition, 936–944 (2017). https://doi.org/10.1109/CVPR.2017.106.",
    "K. He, G. Gkioxari, P. Dollár, and R. Girshick, ‘Mask R-CNN,’ in 2017 IEEE International Conference on Computer Vision, 2980–2988 (2017). https://doi.org/10.1109/ICCV.2017.322.",
    "G.-S. Xia, X. Bai, J. Ding, Z. Zhu, S. Belongie, J. Luo, M. Datcu, M. Pelillo, and L. Zhang, ‘DOTA: a large-scale dataset for object detection in aerial images,’ in Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition, 3974–3983 (2018).",
    "X. Yang, X. Yang, J. Yang, Q. Ming, W. Wang, Q. Tian, and J. Yan, ‘Learning high-precision bounding box for rotated object detection via Kullback–Leibler divergence,’ Advances in Neural Information Processing Systems 34, 18381–18394 (2021).",
    "D. Du, Y. Qi, H. Yu, Y. Yang, K. Duan, G. Li, W. Zhang, Q. Huang, and Q. Tian, ‘The unmanned aerial vehicle benchmark: object detection and tracking,’ in Proceedings of the European Conference on Computer Vision, 370–386 (2018). https://doi.org/10.1007/978-3-030-01249-6_23.",
    "C. Xu, J. Wang, W. Yang, and L. Yu, ‘Dot distance for tiny object detection in aerial images,’ in Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition Workshops, 1192–1201 (2021). https://doi.org/10.1109/CVPRW53098.2021.00130.",
    "S. Zhang, C. Chi, Y. Yao, Z. Lei, and S. Z. Li, ‘Bridging the gap between anchor-based and anchor-free detection via adaptive training sample selection,’ in Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition, 9759–9768 (2020).",
    "Z. Ge, S. Liu, Z. Li, O. Yoshie, and J. Sun, ‘OTA: optimal transport assignment for object detection,’ in Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition, 303–312 (2021). https://doi.org/10.1109/CVPR46437.2021.00037.",
    "X. Yang and J. Yan, ‘Arbitrary-oriented object detection with circular smooth label,’ in Proceedings of the European Conference on Computer Vision, 677–694 (2020). https://doi.org/10.1007/978-3-030-58598-3_40.",
    "X. Yang and J. Yan, ‘Rethinking rotated object detection with Gaussian Wasserstein distance loss,’ in Proceedings of the 38th International Conference on Machine Learning, PMLR 139, 11830–11841 (2021).",
    "J. Ding, N. Xue, Y. Long, G.-S. Xia, and Q. Lu, ‘Learning RoI Transformer for oriented object detection in aerial images,’ in Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition, 2849–2858 (2019). https://doi.org/10.1109/CVPR.2019.00296.",
    "J. Han, J. Ding, N. Xue, and G.-S. Xia, ‘ReDet: a rotation-equivariant detector for aerial object detection,’ in Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition, 2786–2795 (2021). https://doi.org/10.1109/CVPR46437.2021.00281.",
    "X. Xie, G. Cheng, J. Wang, X. Yao, and J. Han, ‘Oriented R-CNN for object detection,’ in Proceedings of the IEEE/CVF International Conference on Computer Vision, 3520–3529 (2021). https://doi.org/10.1109/ICCV48922.2021.00350.",
    "X. Yang, J. Yan, Z. Feng, and T. He, ‘R3Det: refined single-stage detector with feature refinement for rotating object,’ Proceedings of the AAAI Conference on Artificial Intelligence 35(4), 3163–3171 (2021). https://doi.org/10.1609/aaai.v35i4.16426.",
    "B. Li, C. Chen, S. Dong, and J. Qiao, ‘Transmission line detection in aerial images: an instance segmentation approach based on multitask neural networks,’ Signal Processing: Image Communication 96, 116278 (2021). https://doi.org/10.1016/j.image.2021.116278.",
    "B.-E. Benelmostafa and H. Medromi, ‘PowerLine-MTYOLO: a multitask YOLO model for simultaneous cable segmentation and broken strand detection,’ Drones 9(7), 505 (2025). https://doi.org/10.3390/drones9070505.",
    "S. Choudhary, S. Saurav, P. S. Gidde, R. Saini, and S. Singh, ‘LPC-Det: attention-based lightweight object detector for power line component detection in UAV images,’ Computers & Electrical Engineering 126, 110476 (2025). https://doi.org/10.1016/j.compeleceng.2025.110476.",
    "S. Zhang, X. Zhang, W. Ren, L. Shen, and J. Zhang, ‘Shape-aware and feature fused power line detection network,’ Engineering Applications of Artificial Intelligence 170, 113981 (2026). https://doi.org/10.1016/j.engappai.2026.113981.",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_columns(section, count: int, space_twips: int = 360) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        node = cols[0]
    else:
        node = OxmlElement("w:cols")
        sect_pr.append(node)
    node.set(qn("w:num"), str(count))
    node.set(qn("w:space"), str(space_twips))


def set_font(run, name: str = "Times New Roman", size: float | None = None, bold=None, italic=None, color=None) -> None:
    if size is None:
        size = BODY_FONT_SIZE
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_body(doc: Document, text: str, *, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6 if ACTIVE_IVC else 3)
    p.paragraph_format.line_spacing = 1.15 if ACTIVE_IVC else 1.0
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.16)
    set_font(p.add_run(text))


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt((12 if level == 1 else 8) if ACTIVE_IVC else (7 if level == 1 else 4))
    p.paragraph_format.space_after = Pt(6 if ACTIVE_IVC else 2)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(p.add_run(text), size=(12 if level == 1 else 11) if ACTIVE_IVC else (10.5 if level == 1 else 9.5), bold=True)


def add_list_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(text), size=10 if ACTIVE_IVC else 9.5)


def _math_run(parent, text: str) -> None:
    run = OxmlElement("m:r")
    props = OxmlElement("m:rPr")
    style = OxmlElement("m:sty")
    style.set(qn("m:val"), "i")
    props.append(style)
    run.append(props)
    node = OxmlElement("m:t")
    node.text = text.replace("'", "′")
    run.append(node)
    parent.append(run)


def _math_subscript(parent, base: str, sub: str) -> None:
    ssub = OxmlElement("m:sSub")
    base_node = OxmlElement("m:e")
    sub_node = OxmlElement("m:sub")
    _math_run(base_node, base)
    _math_run(sub_node, sub)
    ssub.extend([base_node, sub_node])
    parent.append(ssub)


def _append_linear_math(parent, text: str) -> None:
    """Convert the manuscript's small linear-math subset into editable OMML."""
    import re

    pattern = re.compile(r"([A-Za-zΑ-ωδΔ]+)_\(([^)]+)\)|([A-Za-zΑ-ωδΔ]+)_([A-Za-z0-9]+)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            _math_run(parent, text[cursor : match.start()])
        if match.group(1) is not None:
            _math_subscript(parent, match.group(1), match.group(2))
        else:
            _math_subscript(parent, match.group(3), match.group(4))
        cursor = match.end()
    if cursor < len(text):
        _math_run(parent, text[cursor:])


def add_equation(doc: Document, linear_text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Equation"
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    math_para = OxmlElement("m:oMathPara")
    math = OxmlElement("m:oMath")
    _append_linear_math(math, linear_text)
    math_para.append(math)
    p._p.append(math_para)


def add_equation_plain(doc: Document, display_text: str) -> None:
    """Add a complete editable math line without Word's fragile linear-parser conversion."""
    p = doc.add_paragraph()
    p.style = "Equation Plain"
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    math_para = OxmlElement("m:oMathPara")
    math = OxmlElement("m:oMath")
    _math_run(math, display_text)
    math_para.append(math)
    p._p.append(math_para)


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF2CC")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(text), size=9, bold=True, color=(156, 87, 0))


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Caption"
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.keep_together = True
    set_font(p.add_run(text), size=8.5)


def add_table(doc: Document, headers, rows, widths=None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if ACTIVE_IVC:
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), "9360")
        tbl_w.set(qn("w:type"), "dxa")
        cell_mar = OxmlElement("w:tblCellMar")
        for side, value in (("top", 80), ("left", 120), ("bottom", 80), ("right", 120)):
            node = OxmlElement(f"w:{side}")
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")
            cell_mar.append(node)
        tbl_pr.append(cell_mar)
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(str(text)), size=8, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(str(text)), size=8)
    if widths:
        scale = (6.5 / sum(widths)) if ACTIVE_IVC else 1.0
        for row in table.rows:
            for idx, width in enumerate(widths):
                applied = width * scale
                row.cells[idx].width = Inches(applied)
                tc_w = row.cells[idx]._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    row.cells[idx]._tc.get_or_add_tcPr().append(tc_w)
                tc_w.set(qn("w:w"), str(round(applied * 1440)))
                tc_w.set(qn("w:type"), "dxa")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(BODY_FONT_SIZE)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
    cap = doc.styles["Caption"]
    cap.font.name = "Times New Roman"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cap.font.size = Pt(8.5)
    if "Equation" not in [s.name for s in doc.styles]:
        doc.styles.add_style("Equation", 1)
    eq = doc.styles["Equation"]
    eq.font.name = "Cambria Math"
    eq._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    eq.font.size = Pt(9.5)
    if "Equation Plain" not in [s.name for s in doc.styles]:
        doc.styles.add_style("Equation Plain", 1)
    eq_plain = doc.styles["Equation Plain"]
    eq_plain.font.name = "Cambria Math"
    eq_plain._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    eq_plain.font.size = Pt(9.5)


def add_cover(doc: Document, target: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Inches(1.1)
    set_font(p.add_run("INTERNAL PRE-SUBMISSION BRANCH"), size=20, bold=True, color=(31, 78, 121))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(f"Target: {target['journal']} ({target['short']})"), size=15, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("IVC-specific manuscript rebuilt from the method code and audited evidence; not yet ready for upload"), size=11, italic=True)
    doc.add_paragraph()
    add_note(doc, target["target_note"])
    doc.add_paragraph()
    add_body(doc, "Verified evidence already included: the fixed CA/coarse and CA+Refine validation metrics; exact coarse/identity agreement; 5,260 valid proposals, 3,444 matched proposals, mean matched-IoU change +0.0466, and a 63.41% improvement ratio.", first_line=False)
    add_body(doc, "Mandatory completion items are highlighted in yellow. They are deliberately retained rather than replaced with expected values. Remove this cover page and all yellow notes before submission.", first_line=False)
    doc.add_page_break()


def add_title_block(doc: Document, target: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(target["title"]), size=16 if target["short"] == "IVC" else 15, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Author One, Author Two, and Corresponding Author*"), size=10, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Affiliation, City, Postal Code, Country; *corresponding.author@example.com"), size=9, italic=True)
    add_note(doc, "REQUIRED BEFORE SUBMISSION: replace all author, affiliation, funding, ethics, conflict-of-interest, and data/code availability placeholders. Apply the journal's current disclosure policy at the time of submission.")
    if target["short"] == "IVC":
        add_heading(doc, "Highlights", 1)
        add_list_item(doc, "Finite oriented-regression support is enforced during positive assignment.")
        add_list_item(doc, "Coverage filtering removes geometrically infeasible candidate–target pairs.")
        add_list_item(doc, "Rotated local features refine only the short and long side lengths.")
        add_list_item(doc, "Identity controls isolate refinement from confidence and secondary NMS effects.")
    add_heading(doc, "Abstract", 1)
    add_body(doc, target["abstract"], first_line=False)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    set_font(p.add_run("Keywords: "), size=9.5, bold=True)
    set_font(p.add_run(target["keywords"]), size=9.5)


def add_manuscript(doc: Document, target: dict) -> None:
    add_heading(doc, "1 Introduction")
    add_body(doc, "Low-altitude unmanned aerial vehicles increasingly rely on visual perception for inspection, mapping, emergency response, and navigation near infrastructure. UAV benchmarks identify small objects, camera motion, and changing viewpoints as persistent causes of detection difficulty [22]. Suspended power lines are a particularly demanding obstacle class: they can be difficult to observe before entering the flight corridor, their width may approach the image sampling limit, their length can cross a large portion of a frame, and their apparent orientation changes continuously with the vehicle viewpoint. Reliable detection is therefore a necessary perception function for later risk assessment or path planning, although detection alone does not constitute a complete collision-avoidance system.")
    add_body(doc, "Power-line appearance also varies with illumination, background texture, motion blur, and partial occlusion. Reviews of power-line inspection consistently identify weak appearance, clutter, occlusion, and scale variation as persistent obstacles [2–4]. Localization is particularly brittle when only a few pixels define the short side, because IoU becomes sensitive to small coordinate offsets [23]. The TTPLA dataset provides aerial imagery and annotations for transmission towers and power lines, making it a useful public basis for studying these geometric effects [5].")
    add_body(doc, target["framing"])
    if target["short"] == "IVC":
        add_body(doc, "Recent Image and Vision Computing studies improve low-altitude UAV detection through global feature fusion and self-attention [15], environmentally adaptive context and efficient inference [16], or channel-spatial attention with angle classification in aerial imagery [17]. These methods establish the relevance of UAV and oriented detection to the journal, but their primary interventions are feature enhancement, context adaptation, or angle representation. Our question is narrower and complementary: whether the positive selected on a feature level is geometrically representable, and whether a surviving oriented proposal can be locally rescaled without changing the detector's semantic decision.")
    add_body(doc, "Modern YOLO detectors combine multiscale feature pyramids, task-aligned assignment, and distributional box regression [6–9,18]. ATSS and OTA further demonstrate that the definition of positive samples is itself a central detector design variable [24,25]. In oriented detection, however, a high task-alignment score does not guarantee that a selected feature-level candidate can encode the complete target. The finite support of the distance distribution imposes an explicit geometric range. An elongated object assigned to a shallow level may therefore produce one or more required boundary distances beyond that range, even when the candidate lies inside the rotated ground-truth box. This failure is especially relevant for power lines because the long side dominates the required distance and an off-center candidate amplifies the maximum boundary requirement.")
    add_body(doc, "A second error source appears after coarse detection. At high intersection-over-union (IoU) thresholds, a small error in the short side of a thin box causes a disproportionate overlap loss. Dense per-location residual heads were examined during development but did not provide reliable instance-dependent correction. The final design consequently refines only the proposals that survive nonmaximum suppression (NMS), samples local image features in the proposal coordinate frame, and restricts the update to the two side lengths. This design preserves proposal identity and supports an exact identity control.")
    add_body(doc, "The contributions are threefold. First, a coverage-aware assignment rule makes the finite distance-regression range an explicit condition of positive-sample selection. Second, reg_max=32 is used as the representation support for that rule, rather than claimed as an independent architectural novelty. Third, a proposal-level geometric refiner uses orientation-aligned P2/P3 features to predict bounded short- and long-side scale residuals while leaving center, angle, confidence, and NMS unchanged. The evaluation reports both aggregate accuracy and causal controls that distinguish actual geometric correction from fixed scaling or score recalibration.")

    add_heading(doc, "2 Related Work")
    add_heading(doc, "2.1 Power-line and elongated-object detection", 2)
    add_body(doc, "Early aerial power-line methods combined handcrafted line cues with geometric filtering, whereas more recent systems rely on deep segmentation or object detection [2–4,11]. CableNet uses multitask instance segmentation to preserve individual cable identities [32]; PowerLine-MTYOLO combines cable segmentation with broken-strand detection [33]; LPC-Det targets efficient UAV-based component detection [34]; and SFFPLDN fuses shape-aware multiscale features for visible and infrared line segmentation [35]. Rong et al. introduced a directional processing block to strengthen power-line evidence and used a YOLO-based detector for vegetation-management imagery [1]. These strategies primarily address feature extraction, segmentation, multitask prediction, or efficiency. The present work instead focuses on whether the assigned training candidate can represent the complete oriented target and whether the final proposal has a correct local scale. The directions are potentially complementary but should not be conflated without a controlled combination experiment.")
    add_heading(doc, "2.2 Oriented detection, assignment, and distributed regression", 2)
    add_body(doc, "Task-aligned one-stage detection ranks candidates according to a joint classification-localization score [8], while ATSS and OTA select positives through adaptive statistics or global optimization [24,25]. Generalized focal loss represents each box-side distance with a discrete probability distribution [9]. Aerial benchmarks such as DOTA expose large variation in object orientation and shape [20]. Circular smooth labels, Kullback–Leibler divergence, Gaussian Wasserstein distance, and probabilistic IoU address angle periodicity or rotated-box localization [10,21,26,27]. Existing components improve optimization and angle-aware localization, but the compatibility between a feature level's finite distance support and the complete target has received less direct treatment. Recent remote-sensing work also confirms that large aspect ratios and angle representation remain active concerns [13].")
    add_heading(doc, "2.3 Proposal-level refinement", 2)
    add_body(doc, "Two-stage detectors and instance-segmentation systems demonstrate that proposal-specific features can support a second geometric decision [12,19]. RoI Transformer learns oriented proposal transforms [28], while ReDet and Oriented R-CNN build rotation-aware or oriented proposal pipelines [29,30]. R3Det performs iterative feature refinement in a single-stage detector [31]. A refinement stage, however, can easily duplicate detection, alter ranking, or introduce a second NMS, making attribution ambiguous. The proposed refiner is deliberately narrower: it consumes frozen features and fixed post-NMS proposals, updates only two scale variables, and exposes coarse, identity, and refined modes under the same checkpoint. This restriction is central to the experimental design.")
    if target["short"] == "IVC":
        add_heading(doc, "2.4 Position relative to recent IVC studies", 2)
        comparison_rows = [
            ["Mittal et al. [14]", "Low-altitude UAV survey", "Detection challenges and datasets", "Background and evaluation scope"],
            ["Bai et al. [15]", "UAV small objects", "Self-attention, global fusion, normal-prior assignment", "Closest assignment/ROI context; no finite OBB reachability"],
            ["Sang et al. [16]", "UAV objects", "Adaptive receptive fields and context", "Feature/speed focus; no proposal-identity control"],
            ["Chaurasia and Patro [17]", "Aerial oriented objects", "Channel-spatial attention and CSL", "Oriented detection; not elongated-wire coverage"],
            ["Rong et al. [1]", "Power lines", "Directional filtering and YOLO-OBB", "Closest application domain; feature enhancement focus"],
            ["This work", "Power lines in UAV imagery", "Coverage feasibility and identity-preserving scale refinement", "Representation/assignment/localization focus"],
        ]
        add_table(doc, ["Study", "Task", "Primary intervention", "Relation to this work"], comparison_rows, [1.15, 1.05, 2.0, 2.3])
        add_body(doc, "Cross-paper accuracy values are intentionally not compared because the studies use different datasets, class sets, input sizes, backbones, and evaluation protocols. The comparison instead isolates the algorithmic question addressed by each work.")

    add_heading(doc, "3 Geometric Problem Formulation")
    add_heading(doc, "3.1 Oriented boundary-distance demand", 2)
    add_body(doc, "Let an oriented ground-truth box and a feature-level candidate be defined as")
    add_equation(doc, "B_(j)=(x_(j),y_(j),w_(j),h_(j),θ_(j)),    p_(i)=(x_(i),y_(i))")
    add_body(doc, "After transforming the candidate into the local coordinate system of the target, denote its center offset by (x_f,y_f). The four distances to the local boundaries are w_j/2+x_f, w_j/2−x_f, h_j/2+y_f, and h_j/2−y_f. The largest distance is the minimum scalar range required for complete coverage:")
    add_equation(doc, "D_(req,j)=max(w_j/2+|x_f|,h_j/2+|y_f|)")
    add_body(doc, "For a distributional regressor with reg_max bins, the largest unambiguous distance on level k is")
    add_equation_plain(doc, "Dₘₐₓ,ₖ = 31sₖ,    regₘₐₓ = 32")
    add_body(doc, "Therefore, a candidate can lie inside the ground-truth box while its required distance still exceeds the capacity of the selected feature level. This distinction is the basis of the proposed coverage mask.")
    add_heading(doc, "3.2 Why thin boxes amplify scale errors", 2)
    add_body(doc, "Consider two boxes with aligned centers and angles. If only the short side is perturbed, the relative width error occupies a larger fraction of the area for an intrinsically thin target. Consequently, a modest pixel error can leave AP@0.50 nearly unchanged while sharply reducing AP at stricter IoU thresholds. This observation motivates a local scale refiner but does not justify changing center and angle unless those degrees of freedom are independently supported by oracle or diagnostic evidence.")

    add_heading(doc, "4 Proposed Method")
    add_heading(doc, "4.1 Architecture overview", 2)
    add_body(doc, "The detector follows a YOLO11l-OBB backbone, PAN–FPN neck, and three-scale oriented head. Coverage-aware assignment is active only during training. The box distribution uses reg_max=32. During inference, the standard head produces P3/P4/P5 predictions and NMS returns coarse proposals. A separate refiner extracts rotated regions from P2 and P3 for every retained proposal and writes back only the two side lengths. The detector and refiner are therefore connected at proposal level rather than by adding another dense head.")
    add_heading(doc, "4.2 Coverage-aware positive assignment", 2)
    add_body(doc, "The conventional inside-box candidate mask is intersected with a level-dependent coverage mask. For each candidate-target pair, the coverage mask is one when the required distance does not exceed the feature-level range. The positive candidate set is")
    add_equation(doc, "M_pos=M_in∩M_cov")
    add_body(doc, "The task-aligned ranking is then applied only to geometrically feasible candidates. If the coverage mask removes all candidates for a target, the implementation falls back to the conventional inside-box set. This fallback prevents an empty target assignment and makes the modification conservative: the method changes assignment only when at least one feasible candidate exists.")
    add_heading(doc, "4.3 Regression-range support", 2)
    add_body(doc, "Increasing reg_max from the baseline setting to 32 expands the per-level physical range to 248, 496, and 992 pixels at strides 8, 16, and 32, respectively. The increase is coupled to coverage-aware assignment: it reduces the number of infeasible candidates while preserving an explicit capacity boundary. The larger distribution is treated as a representational support choice and should be evaluated jointly with the assignment rule.")
    add_heading(doc, "4.4 Proposal-level geometric refinement", 2)
    add_body(doc, "For each post-NMS coarse box, the refiner obtains orientation-aligned P2/P3 crops with a 5×24 sampling grid. Lightweight projections, convolution, and a multilayer perceptron fuse the local features. The detector parameters and source features are frozen; only the refiner is optimized. Instead of predicting width and height in image axes, the output is defined on the ordered short and long sides, avoiding the width-height exchange ambiguity of oriented-box parameterizations.")
    add_equation(doc, "δ_(q)=a_(q,−) tanh(r_(q)/a_(q,−)), r_(q)<0")
    add_equation(doc, "δ_(q)=a_(q,+) tanh(r_(q)/a_(q,+)), r_(q)≥0")
    add_equation(doc, "s'=s exp(δ_s),   l'=l exp(δ_l)")
    add_body(doc, "The physical limits are 0.50/0.20 for negative/positive short-side residuals and 0.08/0.08 for the long side. Exact log-scale targets are mapped by the same sign-aware hyperbolic tangent into 80% of these physical ranges, retaining a 20% output margin without a hard-clipping point mass. Matched proposals are optimized with a resolution-weighted smooth-L1 loss (beta=0.05); proposals not selected for geometric supervision receive a zero-residual identity term. The final implementation preserves center, angle, class score, proposal count, and NMS assignment. Identity mode forces both scale residuals to zero and must reproduce the coarse metrics exactly under the same checkpoint and validation configuration.")

    add_heading(doc, "5 Experiments")
    add_heading(doc, "5.1 Dataset and preprocessing", 2)
    add_body(doc, "Experiments use the TTPLA aerial-image dataset [5]. Images are processed at 640×640, and oriented boxes are used for training and validation. The frozen split contains 13,559 training images and 1,695 validation images; the test split has not been used for model selection. The same validation split and preprocessing are used for all numbers reported in this manuscript.")
    add_note(doc, "REQUIRED BEFORE SUBMISSION: insert the audited original-image count, tile-generation rule, class/instance totals, and the final test protocol. Do not infer these values from cache file names.")
    add_heading(doc, "5.2 Implementation and evaluation protocol", 2)
    add_body(doc, "The base detector is YOLO11l-OBB. Coverage-aware assignment uses reg_max=32. The proposal refiner uses P2/P3 features, 32 projected channels, a 5×24 rotated region, and a 128-dimensional hidden representation. The selected refiner was trained for 15 epochs from the frozen CA checkpoint with AdamW, an initial learning rate of 3×10^−4, weight decay 1×10^−4, and three warmup epochs. Checkpoint selection was performed on a holdout protocol, and the reported evaluation was reproduced in FP32 with batch size 8. The random seed is 0.")
    add_body(doc, "Metrics include precision, recall, mAP@0.50, mAP@0.50:0.95, and AP at IoU thresholds 0.75, 0.90, and 0.95. Because only one formal seed is available, no significance claim is made. Differences are reported as controlled single-run effects.")
    add_note(doc, "REQUIRED BEFORE SUBMISSION: recover the unified Baseline and CA training log fields, then report optimizer, batch size, initialization, augmentations, seed, confidence threshold, NMS IoU, max_det, hardware, software versions, and timing warmup consistently.")

    add_heading(doc, "5.3 Main localization results", 2)
    rows = []
    for key in CA:
        rows.append([key, f"{CA[key]:.6f}", f"{REFINE[key]:.6f}", f"{REFINE[key]-CA[key]:+.6f}"])
    add_table(doc, ["Metric", "CA (coarse)", "CA + Refine", "Difference"], rows, [1.20, 0.75, 0.85, 0.75])
    add_body(doc, "The proposal-level refiner improves precision and recall by 0.0895 and 0.0927, respectively. The mAP@0.50:0.95 gain is 0.1084 and the AP@0.75 gain is 0.1190. The improvement becomes smaller at AP@0.90 (+0.0157) and changes sign at AP@0.95 (−0.0018). This threshold profile indicates that the method corrects many moderately inaccurate proposals but does not uniformly improve boxes that are already extremely close to the ground truth.")
    add_note(doc, "REQUIRED BEFORE SUBMISSION: add the original YOLO11l-OBB Baseline row and the Baseline→CA ablation under the same FP32/batch/threshold protocol. The current table establishes only CA→CA+Refine.")

    add_heading(doc, "5.4 Identity and proposal-level diagnostics", 2)
    add_body(doc, "The coarse path bypasses refinement, while identity mode runs the same proposal pathway with zero residuals. Under the same checkpoint, all reported coarse and identity metrics are equal, showing that rotated sampling, feature fusion, and result write-back do not independently change detector output. The normal path therefore differs only through the learned short- and long-side residuals.")
    diag_rows = [
        ["Valid proposals", "5,260"],
        ["Matched proposals", "3,444"],
        ["Mean matched-proposal IoU change", "+0.0466"],
        ["Improved matched proposals", "63.41%"],
        ["Worsened matched proposals", "36.59%"],
        ["Short-side residual boundary hits", "0"],
        ["Long-side residual boundary hits", "0"],
    ]
    add_table(doc, ["Diagnostic", "Value"], diag_rows, [2.4, 1.0])
    add_body(doc, "The positive mean IoU change and the larger improved fraction support instance-dependent correction. Zero boundary hits indicate that the learned residuals are not collapsing to the imposed output limit. These checks do not replace repeated-seed evaluation, but they address two common failure modes: hidden nonidentity operations and uniform fixed-scale shrinkage.")

    if target["short"] == "IVC":
        add_heading(doc, "5.5 Evaluation-path robustness without retraining", 2)
        robustness_rows = [
            ["FP32, batch 8", "0.454137", "0.562504", "+0.108368"],
            ["FP32, batch 1", "0.454151", "0.562460", "+0.108310"],
            ["AMP, batch 8", "0.453562", "0.562017", "+0.108455"],
        ]
        add_table(doc, ["Evaluation path", "CA (coarse)", "CA + Refine", "mAP@0.50:0.95 gain"], robustness_rows, [1.35, 1.0, 1.05, 1.45])
        add_body(doc, "These checks reuse the same trained checkpoint and require no additional optimization. The refinement gain spans only 0.000146 across numerical-precision and batch-size paths, which makes an evaluation-path artifact unlikely. They are reproducibility controls, not independent seeds, and therefore do not support a statistical significance claim.")

    coverage_index = "5.6" if target["short"] == "IVC" else "5.5"
    complexity_index = "5.7" if target["short"] == "IVC" else "5.6"
    add_heading(doc, f"{coverage_index} Coverage mechanism and representation analysis", 2)
    add_body(doc, "The intended mechanism should be verified independently of aggregate AP. For each ground-truth long-side bin, the Baseline and CA models should be compared by (i) the fraction of assigned positives whose required distance exceeds the distributional range and (ii) the normalized P3/P4/P5 assignment proportions. A reduction in overflow together with migration toward a capable level would support the causal explanation. If the measured pattern differs, the interpretation must be revised rather than replaced by an expected diagram.")
    add_note(doc, "REQUIRED BEFORE SUBMISSION: insert the frozen H1/H2 outputs for Baseline and CA. Suggested bins: <100, 100–200, 200–300, 300–500, and >500 pixels. Report counts as well as percentages.")

    add_heading(doc, f"{complexity_index} Complexity, speed, and qualitative behavior", 2)
    add_body(doc, "reg_max=32 increases distributional output channels, while the refiner adds P2/P3 projection, rotated sampling, and a lightweight fusion network. Complexity must therefore be measured end to end with fixed batch size, numerical precision, image size, confidence threshold, NMS threshold, max_det, warmup, and synchronization. Report parameters, FLOPs, peak memory, detector latency, refinement latency, total latency, and frames per second. For the present method, proposal count should also be reported because refiner cost depends on the number of retained boxes.")
    add_note(doc, "REQUIRED BEFORE SUBMISSION: insert verified parameter/FLOP/memory/latency/FPS measurements and an English qualitative panel containing ground truth, Baseline, CA, and CA+Refine under identical thresholds. Include at least one failure case.")

    add_heading(doc, "6 Discussion")
    add_heading(doc, "6.1 Interpretation and scope", 2)
    add_body(doc, "Coverage-aware assignment targets a specific incompatibility between candidate selection and finite distance representation. It is most relevant when long objects can be semantically recognized at a feature level that cannot geometrically cover them. It does not replace feature enhancement for targets that are invisible because of blur, occlusion, or background confusion. Likewise, reg_max=32 should not be presented as a stand-alone algorithmic contribution; its role is to provide sufficient support for the proposed feasibility rule.")
    add_body(doc, "The proposal refiner addresses a different stage. It corrects residual side-length errors after NMS and leaves the detector's semantic decisions unchanged. The strong gains at mAP@0.50:0.95 and AP@0.75, the smaller AP@0.90 increase, and the slight AP@0.95 reduction define a bounded effect: the current model is effective at moving a substantial set of proposals into better overlap regimes, but it is not yet an exact high-IoU optimizer for every instance.")
    add_heading(doc, "6.2 Limitations", 2)
    add_body(doc, "First, numerical reachability does not guarantee adequate semantic context or receptive-field quality. Second, the refiner keeps center and angle fixed, so it cannot repair all sources of oriented-box error. Third, the formal comparison currently contains one random seed. Fourth, the unified Baseline/CA mechanism audit, complexity measurements, and final qualitative figure remain to be inserted under the frozen protocol. These are bounded evidence-completion tasks; they do not justify silently changing the method after the target version has been frozen.")
    if target["short"] == "IVC":
        add_body(doc, "For Image and Vision Computing, the principal remaining risks are evidence completeness and generality. The comparison still lacks a unified original-baseline row, the coverage-overflow mechanism statistics, end-to-end computational cost, repeated seeds, and a frozen test-set evaluation. These items must be completed without changing the already stated method. Until then, the manuscript is a structurally complete pre-submission draft rather than a submission-ready claim package.")
    elif target["short"] == "US":
        add_body(doc, "For UAV deployment, the present results validate only an image-based perception module. They do not yet demonstrate metric wire distance, closed-loop trajectory generation, collision avoidance, or robustness under flight dynamics. Complete detector-plus-refiner latency must be measured, including NMS and rotated proposal sampling; an edge-computing or onboard benchmark would be needed before describing the system as real time or deployment ready.")
    elif target["short"] == "JARS":
        add_body(doc, "For remote-sensing generalization, the present evidence is limited to the TTPLA low-altitude aerial domain. Evaluation on another large-aspect-ratio oriented remote-sensing benchmark would strengthen claims beyond power-line inspection. Until such evidence is available, the conclusion is restricted to the reported aerial dataset and acquisition conditions.")
    else:
        add_body(doc, "For imaging-system deployment, latency and memory must be measured on the intended hardware. The method should not be described as real time until complete detector-plus-refiner throughput, including NMS and proposal sampling, has been verified.")

    add_heading(doc, "7 Conclusion")
    add_body(doc, "This work formulates power-line obstacle detection for low-altitude UAV perception as a compatibility problem between feature-level assignment, finite distributional regression, and proposal-scale precision. Coverage-aware assignment filters candidates that cannot represent the complete rotated target, and reg_max=32 supplies a larger but still explicit distance range. A proposal-level refiner then uses orientation-aligned local features to update only short and long sides. In the fixed 640-pixel FP32 evaluation, CA+Refine improves mAP@0.50:0.95 from 0.4541 to 0.5625 and AP@0.75 from 0.4398 to 0.5588. Identity and proposal-level diagnostics support geometric correction as the source of the gain. The evidence is positive but deliberately bounded by the small AP@0.90 gain, the slight AP@0.95 decrease, and the current single-seed setting. The resulting oriented detections provide an image-space perception output for subsequent UAV safety reasoning; integration with ranging, navigation, and flight control remains future work.")

    add_heading(doc, "Disclosures")
    add_body(doc, "Conflict of interest: [TO BE COMPLETED].", first_line=False)
    add_body(doc, "Funding: [TO BE COMPLETED].", first_line=False)
    add_body(doc, "Data availability: TTPLA is publicly available; the exact processed split and preprocessing statement will be added before submission.", first_line=False)
    add_body(doc, "Code availability: [TO BE COMPLETED according to the selected journal and institutional policy].", first_line=False)
    add_body(doc, "Author contributions: [TO BE COMPLETED].", first_line=False)
    if target["short"] == "IVC":
        add_heading(doc, "Declaration of generative AI and AI-assisted technologies in the writing process")
        add_body(doc, "During preparation of this work, the authors used OpenAI Codex to assist with manuscript drafting, language editing, figure preparation, and document formatting. The authors reviewed and edited the resulting content and take full responsibility for the content of the publication. The final wording must be checked against the journal's policy in force at submission.", first_line=False)

    add_heading(doc, "References")
    for idx, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(1.5)
        set_font(p.add_run(f"[{idx}] {ref}"), size=8)


def add_figure_plates(doc: Document, target: dict) -> None:
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_columns(section, 1)
    add_heading(doc, "Figure Plates for Internal Layout Review")
    if target["short"] == "IVC":
        add_note(doc, "Editable English figure sources are stored beside the raster images. The figures contain no internal experimental version labels. Formula objects in Figs. 1–2 are editable Office Math objects in the source PPTX files.")
        figure_root = IVC_FIGURES
        figures = [
            ("fig1_ivc_method.png", "Fig. 1. Architecture of coverage-aware YOLO11-OBB and identity-preserving proposal-level geometric refinement."),
            ("fig2_ivc_reachability.png", "Fig. 2. Candidate-local boundary demand and feature-level coverage feasibility. The numerical example illustrates the mechanism and is not a measured training result."),
            ("fig3_ivc_refinement_results.png", "Fig. 3. Checkpoint selection on the training holdout and independent FP32 validation of proposal-level refinement."),
        ]
    else:
        add_note(doc, "The source figures below are the current verified Chinese assets. Before journal submission, export English-only versions with the same equations and values; do not translate by raster overlay.")
        figure_root = SOURCE_FIGURES
        figures = [
            ("ca_refine_architecture_redesign.png", "Fig. 1 Overall architecture of coverage-aware YOLO11-OBB and proposal-level geometric refinement."),
            ("fig2_geometric_reachability.png", "Fig. 2 Candidate-local boundary demand and level-dependent coverage feasibility."),
            ("fig4_refine_v311_curve.png", "Fig. 3 Refiner checkpoint selection and independent FP32 localization results."),
        ]
    for index, (filename, caption) in enumerate(figures):
        if index:
            doc.add_page_break()
        path = figure_root / filename
        if path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(path), width=Inches(6.85))
            add_caption(doc, caption)
        else:
            add_note(doc, f"MISSING FIGURE: {path}")


def add_page_numbers(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_font(run, size=8)


def build(target_key: str, out_dir: Path) -> Path:
    global ACTIVE_IVC, BODY_FONT_SIZE
    target = TARGETS[target_key]
    ACTIVE_IVC = target["short"] == "IVC"
    BODY_FONT_SIZE = 10.5 if ACTIVE_IVC else 9.5
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    margin = 1.0 if ACTIVE_IVC else 0.62
    side_margin = 1.0 if ACTIVE_IVC else 0.68
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(side_margin)
    section.right_margin = Inches(side_margin)
    section.header_distance = Inches(0.492 if ACTIVE_IVC else 0.25)
    section.footer_distance = Inches(0.492 if ACTIVE_IVC else 0.28)
    set_columns(section, 1)
    add_cover(doc, target)
    add_title_block(doc, target)
    body_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    # Keep the editable review manuscript single-column. SPIE applies the
    # final two-column production layout after acceptance; this also avoids
    # unstable pagination in long Word drafts with full-width figure plates.
    set_columns(body_section, 1)
    body_section.top_margin = Inches(margin)
    body_section.bottom_margin = Inches(margin)
    body_section.left_margin = Inches(side_margin)
    body_section.right_margin = Inches(side_margin)
    add_manuscript(doc, target)
    add_figure_plates(doc, target)
    # All section footers remain linked to the first section. Adding the PAGE
    # field once prevents the same field from being duplicated as 111/222/...
    # when the manuscript and figure-plate sections are created.
    add_page_numbers(doc.sections[0])
    props = doc.core_properties
    props.title = target["title"]
    props.subject = f"Internal pre-submission branch for {target['journal']}"
    props.keywords = target["keywords"]
    props.comments = "Contains highlighted completion notes; not ready for upload."
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"{target['file_stem']}_投稿预备分支_英文版.docx"
    doc.save(out)
    return out


def main() -> None:
    parser = argparse.ArgumentParser(description="Build target-specific English pre-submission branches for UAV, remote-sensing, and imaging journals.")
    parser.add_argument("--output-dir", type=Path, default=ROOT / "mydocs" / "创新点一" / "投稿版本")
    parser.add_argument("--target", choices=[*TARGETS, "all"], default="all")
    args = parser.parse_args()
    keys = TARGETS if args.target == "all" else [args.target]
    for key in keys:
        print(build(key, args.output_dir))


if __name__ == "__main__":
    main()
